#!/usr/bin/env python3
"""
Economic News Aggregator
========================
Reads RSS feeds from "RSS Feeds.docx" and keywords from "Keywords.docx",
filters articles from the last 24 hours, scores them by keyword relevance,
generates AI summaries via a locally running Ollama (mistral) model, and
writes a styled, searchable, mobile-responsive HTML digest.

Usage:
    python aggregator.py

Requirements:
    pip install feedparser requests python-docx beautifulsoup4 scikit-learn
    Ollama running locally with the mistral model pulled:
        ollama pull mistral
        ollama serve
"""

import sys
import os
import re
import json
import html as html_module
from datetime import datetime, timedelta, timezone
from concurrent.futures import ThreadPoolExecutor, as_completed

# Third-party (install with pip if missing)
try:
    import feedparser
except ImportError:
    sys.exit("Missing package: run  pip install feedparser")

try:
    import requests
except ImportError:
    sys.exit("Missing package: run  pip install requests")

try:
    from docx import Document
except ImportError:
    sys.exit("Missing package: run  pip install python-docx")

try:
    from bs4 import BeautifulSoup
except ImportError:
    sys.exit("Missing package: run  pip install beautifulsoup4")

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity as cos_sim
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False
    print("Warning: scikit-learn not installed — cosine-similarity clustering disabled. "
          "Run: pip install scikit-learn")


# =============================================================================
# CONFIGURATION — adjust these if needed
# =============================================================================

BASE_DIR             = os.path.dirname(os.path.abspath(__file__))
FEEDS_DOCX           = os.path.join(BASE_DIR, "RSS Feeds.docx")
KEYWORDS_DOCX        = os.path.join(BASE_DIR, "Keywords.docx")
OUTPUT_HTML          = os.path.join(BASE_DIR, "digest.html")
KEYWORD_HISTORY_JSON = os.path.join(BASE_DIR, "keyword_history.json")

# ── Email settings ────────────────────────────────────────────────────────────
# To enable email: create a Gmail App Password at
# https://myaccount.google.com/apppasswords (requires 2-Step Verification)
# then paste it below.
EMAIL_TO           = '@gmail.com'
EMAIL_FROM         = '@gmail.com'
EMAIL_APP_PASSWORD = ''

# Maps the region strings that come from RSS Feeds.docx section headers
# to the six display region slugs used in the digest nav and badges.
REGION_SLUG_MAP = {
    'IRELAND':                            'ireland',
    'UNITED KINGDOM':                     'uk',
    'EUROPE — WESTERN':              'europe',
    'EUROPE — EASTERN & CENTRAL':    'europe',
    'RUSSIA & FORMER SOVIET STATES':      'europe',
    'UNITED STATES':                      'us',
    'CANADA':                             'canada',
    'LATIN AMERICA':                      'latam',
    'CHINA & GREATER CHINA':              'asia-pacific',
    'INDIA':                              'asia-pacific',
    'JAPAN & SOUTH KOREA':                'asia-pacific',
    'SOUTHEAST ASIA':                     'asia-pacific',
    'AUSTRALIA & NEW ZEALAND':            'asia-pacific',
    'AFRICA':                             'africa-me',
    'MIDDLE EAST':                        'africa-me',
    'ACADEMIC & PREPRINT':                'global',
    'DATA & INDICATOR RELEASES':          'global',
    'INDEPENDENT ECONOMISTS & SUBSTACKS': 'global',
    'INTERNATIONAL & MULTILATERAL':       'global',
    'INTERNATIONAL THINK TANKS':          'global',
}

HOURS_BACK     = 24   # Only articles published in the last N hours
TOP_N          = 5    # Number of highlighted top articles
MAX_WORKERS    = 20   # How many feeds to fetch in parallel
FETCH_TIMEOUT  = 12   # Seconds before giving up on a single feed
OLLAMA_URL     = "http://localhost:11434/api/generate"
OLLAMA_MODEL   = "mistral"
OLLAMA_TIMEOUT = 90   # Seconds to wait for each summary
MIN_SUBCATS    = 15   # Category needs at least this many articles to get sub-sections

# Priority order for deciding which theme "owns" an article that matches several.
# Most-specific first so Irish stories aren't swallowed by Global Economy, etc.
THEME_PRIORITY = [
    'IRELAND',
    'HOUSING & PROPERTY',
    'DEMOGRAPHICS & MIGRATION',
    'INEQUALITY & LABOUR',
    'EUROPE & ECB',
    'TRADE & GLOBALISATION',
    'MACROECONOMICS',
    'GLOBAL ECONOMY',
]

# Display order for the category nav tabs and home-page grid.
CATEGORY_DISPLAY_ORDER = [
    'IRELAND',
    'HOUSING & PROPERTY',
    'MACROECONOMICS',
    'GLOBAL ECONOMY',
    'TRADE & GLOBALISATION',
    'INEQUALITY & LABOUR',
    'DEMOGRAPHICS & MIGRATION',
    'General',
    'EUROPE & ECB',
]

# Common English stopwords stripped when comparing article titles for clustering.
_STOPWORDS = {
    'a','an','the','in','on','at','of','to','for','is','are','was','were',
    'be','been','being','have','has','had','do','does','did','will','would',
    'could','should','may','might','shall','can','and','or','but','not',
    'with','from','by','as','its','it','this','that','their','they',
    'he','she','we','i','you','up','out','than','more','over','new',
    'about','after','into','how','s','us',
}

# Keyword clusters used to split large categories (≥ MIN_SUBCATS) into sub-sections.
# Each entry: (sub-category display name, list of lowercase keyword fragments).
# An article is placed in the FIRST sub-category whose keywords appear in its
# title + summary + matched keywords blob.
SUBCATEGORY_RULES = {
    'GLOBAL ECONOMY': [
        ('US Economy',          ['united states', 'us economy', 'federal reserve', 'fed ', 'trump', 'wall street', 'washington', 'american economy', 'us gdp', 'us jobs', 'us dollar']),
        ('China Economy',       ['china', 'chinese', 'beijing', 'yuan', 'xi jinping', 'prc ', 'hong kong economy', 'taiwan economy']),
        ('European Economy',    ['eurozone', 'euro area', 'germany', 'france', 'italy', 'spain', 'eu economy', 'european economy', 'ecb rate']),
        ('UK Economy',          ['united kingdom', 'british economy', 'uk economy', 'bank of england', 'sterling', 'british gdp', 'uk jobs']),
        ('Energy & Commodities',['oil price', 'opec', 'energy crisis', 'natural gas', 'crude oil', 'commodity', 'coal price', 'oil shock', 'gas price', 'lng']),
        ('Trade & Tariffs',     ['tariff', 'trade war', 'trade deficit', 'trade surplus', 'export ban', 'import duty', 'supply chain', 'wto', 'customs duty']),
        ('Emerging Markets',    ['emerging market', 'india economy', 'brazil economy', 'indonesia', 'vietnam', 'developing economies', 'mexico economy', 'africa economy', 'southeast asia']),
        ('Financial Markets',   ['stock market', 'bond yield', 'equity market', 'market crash', 'market rally', 'financial crisis', 'banking sector', 'interest rate']),
    ],
    'IRELAND': [
        ('Housing & Property',  ['housing', 'rent ', 'property', 'mortgage', 'landlord', 'tenant', 'planning permission', 'construction', 'house price']),
        ('Labour & Employment', ['employment', 'jobs', 'unemployment', 'workers', 'wage', 'labour market', 'hiring', 'layoffs', 'redundanc']),
        ('Government & Fiscal', ['budget', 'government', 'fiscal', 'taxation', 'public spending', 'national debt', 'oireachtas', 'dail', 'minister']),
        ('Business & FDI',      ['business', 'investment', 'fdi', 'multinational', 'startup', 'enterprise', 'pharma', 'tech company', 'corporation']),
        ('Economy & Growth',    ['gdp', 'growth', 'inflation', 'recession', 'economic', 'output', 'gnp', 'central bank of ireland']),
    ],
    'MACROECONOMICS': [
        ('Inflation & Prices',                ['inflation', 'cpi', 'price level', 'deflation', 'cost of living', 'price rise', 'purchasing power', 'price index']),
        ('Interest Rates & Monetary Policy',  ['interest rate', 'central bank', 'monetary policy', 'rate cut', 'rate hike', 'quantitative easing', 'base rate', 'policy rate']),
        ('Growth & GDP',                      ['gdp', 'economic growth', 'recession', 'contraction', 'expansion', 'output gap', 'gdp growth']),
        ('Labour & Employment',               ['unemployment rate', 'labour market', 'payroll', 'jobless', 'employment rate', 'job market', 'nonfarm']),
    ],
    'TRADE & GLOBALISATION': [
        ('Tariffs & Trade Wars',    ['tariff', 'trade war', 'trade dispute', 'protectionism', 'customs duty', 'retaliation', 'trade barrier', 'import tariff']),
        ('Supply Chains',           ['supply chain', 'logistics', 'shipping', 'port ', 'manufacturing', 'reshoring', 'offshoring', 'nearshoring']),
        ('Trade Agreements',        ['trade agreement', 'free trade', 'trade deal', 'bilateral', 'multilateral', 'wto ruling', 'trade negotiations']),
        ('Globalisation',           ['globalisation', 'globalization', 'deglobalisation', 'decoupling', 'world trade', 'international trade']),
    ],
    'EUROPE & ECB': [
        ('ECB & Monetary Policy',   ['ecb', 'european central bank', 'lagarde', 'rate decision', 'monetary policy', 'quantitative easing', 'deposit rate']),
        ('EU Policy',               ['european commission', 'european parliament', 'european council', 'eu regulation', 'brussels', 'von der leyen', 'eu law']),
        ('Eurozone Economy',        ['eurozone', 'euro area', 'single currency', 'stability pact', 'fiscal rules', 'sovereign debt']),
        ('EU Countries',            ['germany', 'france', 'italy', 'spain', 'netherlands', 'poland', 'austria', 'greece', 'belgium', 'czech', 'hungarian']),
    ],
    'HOUSING & PROPERTY': [
        ('Prices & Affordability',  ['house price', 'property price', 'affordability', 'market value', 'sale price', 'price rise', 'home price']),
        ('Rental Market',           ['rent ', 'rental', 'landlord', 'tenant', 'eviction', 'lease', 'renter', 'rental income']),
        ('Planning & Development',  ['planning', 'development', 'construction', 'build', 'zoning', 'density', 'planning permission', 'developer']),
        ('Policy & Schemes',        ['government scheme', 'help to buy', 'social housing', 'affordable housing', 'subsidy', 'housing policy']),
    ],
    'INEQUALITY & LABOUR': [
        ('Wages & Income',          ['wage', 'salary', 'pay ', 'minimum wage', 'pay gap', 'income inequality', 'earnings', 'compensation']),
        ('Workers & Unions',        ['trade union', 'union ', 'strike', 'collective bargaining', 'workers rights', 'industrial action', 'labour dispute']),
        ('Poverty & Welfare',       ['poverty', 'welfare', 'social protection', 'inequality', 'deprivation', 'low income', 'benefit']),
        ('Future of Work',          ['automation', 'ai jobs', 'artificial intelligence', 'robot', 'gig economy', 'remote work', 'platform work', 'future of work']),
    ],
    'DEMOGRAPHICS & MIGRATION': [
        ('Immigration & Asylum',    ['immigration', 'immigrant', 'migrant', 'asylum', 'refugee', 'border control', 'visa', 'deportation']),
        ('Population & Ageing',     ['ageing', 'aging', 'birth rate', 'fertility rate', 'population growth', 'demographic', 'elderly', 'pension age']),
        ('Emigration & Diaspora',   ['emigration', 'diaspora', 'brain drain', 'emigrant', 'return migration']),
    ],
}


# =============================================================================
# STEP 1 — READ SOURCES FROM RSS FEEDS.DOCX
# =============================================================================
# The docx uses this format:
#   # =====  REGION NAME  =====    ← major region header
#   # --- Sub-section ---          ← sub-section (we ignore these for grouping)
#   Label | https://feed.url       ← actual feed lines
#
# Lines starting with # are comments/headers; blank lines are ignored.

def read_sources():
    """Return a list of dicts: {label, url, region}."""
    doc = Document(FEEDS_DOCX)
    sources = []
    current_region = "International"

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # '## ' prefix marks a disabled feed — kept for reference, not fetched
        if text.startswith('##'):
            continue

        if text.startswith('#'):
            inner = text.lstrip('#').strip()
            # Skip pure decoration (===, ---, empty after stripping)
            if not inner or re.match(r'^[=\-\s]+$', inner):
                continue
            # Sub-section headers like "--- Irish News ---" — skip, keep region
            if inner.startswith('---') or re.match(r'^-+\s', inner):
                continue
            # Skip lines that are clearly instructions, not region names
            skip_words = ('Format:', 'To add', 'To remove', 'Total sources',
                          'Last updated', 'Or Google', 'END OF', 'News Aggregator',
                          'find RSS', 'adding /feed')
            if any(w in inner for w in skip_words):
                continue
            # What remains is a region name
            current_region = inner
            continue

        # Feed line: must contain "|" and a URL
        if '|' in text:
            parts = text.split('|', 1)
            if len(parts) == 2:
                label = parts[0].strip()
                # Strip ALL unicode whitespace including zero-width spaces,
                # non-breaking spaces, etc. that Word sometimes inserts
                url = re.sub(r'[\s\u00a0\u200b\u200c\u200d\ufeff]+', '', parts[1].strip())
                if url.startswith('http'):
                    sources.append({'label': label, 'url': url, 'region': current_region})

    return sources


def count_disabled_feeds():
    """Return the count of ## -prefixed feed lines in RSS Feeds.docx."""
    doc   = Document(FEEDS_DOCX)
    total = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith('##') and '|' in text:
            parts = text.lstrip('#').split('|', 1)
            if len(parts) == 2 and parts[1].strip().startswith('http'):
                total += 1
    return total


# =============================================================================
# STEP 2 — READ KEYWORDS FROM KEYWORDS.DOCX
# =============================================================================
# Format:
#   # --- THEME NAME ---   ← section header
#   keyword or phrase      ← one per line; case-insensitive matching later
#   # disabled keyword     ← lines starting with # are skipped

def read_keywords():
    """Return a list of dicts: {keyword, theme}."""
    doc = Document(KEYWORDS_DOCX)
    keywords = []
    current_theme = "General"

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text.startswith('#'):
            inner = text.lstrip('#').strip()
            # Theme headers look like "--- HOUSING & PROPERTY ---"
            if '---' in inner:
                theme = inner.replace('---', '').strip()
                if theme:
                    current_theme = theme
            continue

        # It's a keyword/phrase line
        kw = text.lower()
        if kw:
            keywords.append({'keyword': kw, 'theme': current_theme})

    return keywords


# =============================================================================
# STEP 3 — FETCH ALL RSS FEEDS IN PARALLEL
# =============================================================================

def fetch_one_feed(source):
    """
    Fetch a single feed URL and return (label, region, entries, error_type).
    error_type is None on success, 'routine' for expected 403/404 failures,
    or 'unexpected' for genuine problems worth reporting.
    """
    label  = source['label']
    url    = source['url']
    region = source['region']
    try:
        resp = requests.get(
            url,
            timeout=FETCH_TIMEOUT,
            headers={'User-Agent': 'Mozilla/5.0 (compatible; EconNewsAggregator/1.0)'}
        )
        resp.raise_for_status()
        feed = feedparser.parse(resp.content)
        return label, region, feed.entries, None
    except requests.exceptions.HTTPError as exc:
        code = exc.response.status_code if exc.response is not None else 0
        # 403 = site blocks scrapers, 404 = feed moved/deleted — both routine
        if code in (403, 404, 410):
            return label, region, [], 'routine'
        return label, region, [], f'HTTP {code}'
    except requests.exceptions.Timeout:
        return label, region, [], 'timeout'
    except requests.exceptions.ConnectionError as exc:
        msg = str(exc)
        if 'NameResolutionError' in msg or 'getaddrinfo failed' in msg:
            return label, region, [], 'dns'
        if 'SSLError' in msg or 'SSL' in msg:
            return label, region, [], 'ssl'
        return label, region, [], 'connection'
    except Exception as exc:
        return label, region, [], f'error: {exc}'


def fetch_all_feeds(sources):
    """
    Fetch all feeds using a thread pool for speed.
    Returns a flat list of feed entries, each tagged with _source_label
    and _source_region. Prints a clean summary of failures instead of
    one noisy line per broken feed.
    """
    print(f"Fetching {len(sources)} feeds with up to {MAX_WORKERS} parallel connections...")
    all_entries = []

    # Counters for the summary line
    failed = {'routine': 0, 'timeout': 0, 'dns': 0, 'ssl': 0, 'connection': 0}
    unexpected = []   # (label, error_type) for non-routine failures

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as pool:
        futures = {pool.submit(fetch_one_feed, s): s for s in sources}
        completed = 0
        for future in as_completed(futures):
            completed += 1
            label, region, entries, err = future.result()

            if err is None:
                for entry in entries:
                    entry['_source_label']  = label
                    entry['_source_region'] = region
                all_entries.extend(entries)
            elif err == 'routine':
                failed['routine'] += 1
            elif err in failed:
                failed[err] += 1
                unexpected.append((label, err))
            else:
                # Miscellaneous unexpected error
                failed['connection'] = failed.get('connection', 0) + 1
                unexpected.append((label, err))

            if completed % 25 == 0 or completed == len(sources):
                print(f"  Progress: {completed}/{len(sources)} feeds fetched "
                      f"({len(all_entries)} raw entries so far)")

    print(f"Total raw entries across all feeds: {len(all_entries)}")

    # Print a clean summary instead of one line per failure
    routine_count = failed.pop('routine', 0)
    other_counts  = {k: v for k, v in failed.items() if v > 0}
    if routine_count:
        print(f"  Skipped (blocked/moved feeds): {routine_count}")
    if other_counts:
        summary = ', '.join(f"{v} {k}" for k, v in other_counts.items())
        print(f"  Skipped (network issues): {summary}")
    if unexpected:
        print("  Feeds with unexpected errors (may be worth investigating):")
        for lbl, err in unexpected:
            print(f"    • {lbl}: {err}")

    return all_entries


# =============================================================================
# STEP 4 — FILTER TO LAST 24 HOURS
# =============================================================================

def entry_published_time(entry):
    """
    Extract the publication time from a feed entry as a timezone-aware datetime.
    Returns None if no valid time can be found.
    feedparser gives us time as a struct_time tuple in 'published_parsed' or
    'updated_parsed'. We convert that to a UTC datetime.
    """
    for field in ('published_parsed', 'updated_parsed'):
        t = entry.get(field)
        if t:
            try:
                return datetime(*t[:6], tzinfo=timezone.utc)
            except Exception:
                pass
    return None


def extract_image(entry):
    """Return the best available thumbnail URL from a feed entry, or None."""
    # 1. media:thumbnail (most common in news feeds)
    thumbs = getattr(entry, 'media_thumbnail', None)
    if thumbs:
        url = thumbs[0].get('url', '')
        if url:
            return url
    # 2. media:content with an image type
    mc_list = getattr(entry, 'media_content', None)
    if mc_list:
        for mc in mc_list:
            url = mc.get('url', '')
            if url and 'image' in mc.get('type', 'image'):
                return url
    # 3. enclosures
    for enc in getattr(entry, 'enclosures', []):
        url = enc.get('href') or enc.get('url', '')
        if url and 'image' in enc.get('type', ''):
            return url
    # 4. First <img> in the HTML content/summary
    html_src = ''
    if entry.get('content'):
        html_src = entry['content'][0].get('value', '')
    if not html_src:
        html_src = entry.get('summary', '') or ''
    if html_src:
        soup = BeautifulSoup(html_src, 'html.parser')
        img = soup.find('img')
        if img:
            src = img.get('src', '')
            if src and src.startswith('http'):
                return src
    return None


def filter_recent(entries):
    """Keep only entries published within the last HOURS_BACK hours."""
    cutoff = datetime.now(timezone.utc) - timedelta(hours=HOURS_BACK)
    recent = [e for e in entries if (t := entry_published_time(e)) and t >= cutoff]
    print(f"Articles from the last {HOURS_BACK} hours: {len(recent)}")
    return recent


# =============================================================================
# STEP 5 — SCORE ARTICLES BY KEYWORD MATCHES
# =============================================================================

def strip_html(raw_html):
    """Remove HTML tags and return plain text."""
    if not raw_html:
        return ''
    soup = BeautifulSoup(raw_html, 'html.parser')
    return soup.get_text(separator=' ')


def score_entry(entry, keywords):
    """
    Build a combined text blob from the title + summary + full content,
    then count how many keywords appear in it.

    Returns:
        score          — integer count of keyword matches
        matched_themes — set of theme names that were matched
        matched_kws    — list of the actual keywords that matched (for display)
    """
    title   = entry.get('title', '')
    summary = strip_html(entry.get('summary', '') or entry.get('description', ''))
    content = ''
    if entry.get('content'):
        content = strip_html(' '.join(c.get('value', '') for c in entry['content']))

    # Single lowercase blob to search
    blob = (title + ' ' + summary + ' ' + content).lower()
    # Remove leftover HTML noise
    blob = re.sub(r'<[^>]+>', ' ', blob)

    matched_themes = set()
    matched_kws    = []

    for kw_dict in keywords:
        word = kw_dict['keyword']
        if ' ' in word:
            # Multi-word phrase: simple substring match
            if word in blob:
                matched_themes.add(kw_dict['theme'])
                matched_kws.append(word)
        else:
            # Single word: use word boundaries so "rate" doesn't match "moderate"
            if re.search(r'\b' + re.escape(word) + r'\b', blob):
                matched_themes.add(kw_dict['theme'])
                matched_kws.append(word)

    return len(matched_kws), matched_themes, matched_kws


# =============================================================================
# STEP 6 — DEDUPLICATE
# =============================================================================

def normalize_title(title):
    """Lowercase, strip punctuation — used as a deduplication fingerprint."""
    t = title.lower()
    t = re.sub(r'[^a-z0-9 ]', '', t)
    t = re.sub(r'\s+', ' ', t).strip()
    return t


def deduplicate(articles):
    """
    Remove articles whose titles are essentially identical.
    The same story often appears across multiple syndicated feeds.
    We use the first 70 characters of the normalized title as a fingerprint.
    """
    seen    = set()
    unique  = []
    for art in articles:
        fp = normalize_title(art['title'])[:70]
        if fp not in seen:
            seen.add(fp)
            unique.append(art)
    print(f"After removing duplicates: {len(unique)} articles")
    return unique


# =============================================================================
# STEP 6.5 — CLUSTER RELATED STORIES
# =============================================================================

def _sig_words(title):
    """Significant words in a title: lowercase, no punctuation, no stopwords, len > 2."""
    words = re.sub(r'[^\w\s]', ' ', title.lower()).split()
    return set(w for w in words if w not in _STOPWORDS and len(w) > 2)


def cluster_articles(articles):
    """
    Group articles that cover the same underlying event.

    Two articles are in the same cluster when their titles share ≥ 3 significant
    words OR their AI summaries have cosine TF-IDF similarity > 0.65 (requires
    scikit-learn).  Within each cluster the article with the highest relevance
    score becomes the 'lead'; the others become 'also covering' entries attached
    to the lead as a list of {title, url, source} dicts.

    Solo articles (no cluster match) receive an empty 'also' list and are
    otherwise unchanged.  The returned list is score-sorted leads only.
    """
    n = len(articles)
    if n == 0:
        return articles

    # ── TF-IDF cosine similarity matrix ──────────────────────────────────────
    sim_matrix = None
    if SKLEARN_AVAILABLE and n >= 2:
        texts = [art.get('summary') or art.get('title', '') for art in articles]
        try:
            vec    = TfidfVectorizer(stop_words='english', max_features=5000)
            tfidf  = vec.fit_transform(texts)
            sim_matrix = cos_sim(tfidf)          # shape (n, n), float32
        except Exception:
            pass  # fall back to title-word matching only

    # ── Precompute significant title words ───────────────────────────────────
    sig = [_sig_words(art['title']) for art in articles]

    # ── Union-Find ────────────────────────────────────────────────────────────
    parent = list(range(n))

    def find(x):
        while parent[x] != x:
            parent[x] = parent[parent[x]]
            x = parent[x]
        return x

    def union(x, y):
        parent[find(x)] = find(y)

    for i in range(n):
        for j in range(i + 1, n):
            title_match = len(sig[i] & sig[j]) >= 3
            cos_match   = (sim_matrix is not None and float(sim_matrix[i, j]) > 0.65)
            if title_match or cos_match:
                union(i, j)

    # ── Group by cluster root ─────────────────────────────────────────────────
    clusters: dict = {}
    for i in range(n):
        clusters.setdefault(find(i), []).append(i)

    # ── Build output: one lead per cluster ────────────────────────────────────
    result = []
    for members in clusters.values():
        if len(members) == 1:
            articles[members[0]]['also'] = []
            articles[members[0]]['perspectives'] = False
            result.append(articles[members[0]])
        else:
            # Lead = highest score; tie-break by original (score-sorted) index
            members.sort(key=lambda i: (-articles[i]['score'], i))
            lead = articles[members[0]]

            # Check if the cluster spans 3+ distinct region_slugs
            cluster_slugs = {articles[k].get('region_slug', 'global') for k in members}
            is_perspectives = len(cluster_slugs) >= 3

            if is_perspectives:
                # One entry per region_slug (highest-scored = first encountered
                # because members is already score-sorted)
                seen_regions: set = set()
                also_list = []
                for k in members[1:]:
                    rslug = articles[k].get('region_slug', 'global')
                    if rslug not in seen_regions:
                        seen_regions.add(rslug)
                        also_list.append({
                            'title':       articles[k]['title'],
                            'url':         articles[k]['link'],
                            'source':      articles[k]['source'],
                            'region':      articles[k].get('region', ''),
                            'region_slug': rslug,
                            'summary':     articles[k].get('summary', ''),
                        })
                lead['also'] = also_list
                lead['perspectives'] = True
            else:
                lead['also'] = [
                    {'title':  articles[k]['title'],
                     'url':    articles[k]['link'],
                     'source': articles[k]['source']}
                    for k in members[1:]
                ]
                lead['perspectives'] = False

            result.append(lead)

    # Preserve score-descending order among leads
    result.sort(key=lambda a: -a['score'])

    clustered = sum(len(a['also']) for a in result)
    if clustered:
        print(f"Clustering: {clustered} article(s) folded into {len(result)} leads "
              f"(page reduced from {n} to {len(result)} stories)")
    return result


def mark_underreported(articles):
    """
    Flag high-relevance solo stories that appear in only one geographic source.

    Criteria (all three must hold):
      • score >= 1.5 × median score across all articles
      • also == []  (no cluster partner — single-source coverage)
      • region_slug not 'global' (real geographic source, not think tank / academic)

    Sets art['underreported'] = True|False on every article in-place and
    returns the count of flagged articles.
    """
    if not articles:
        return 0
    scores = sorted(a['score'] for a in articles)
    mid    = len(scores) // 2
    median = (scores[mid] if len(scores) % 2 else (scores[mid - 1] + scores[mid]) / 2)
    threshold = 1.5 * median
    count = 0
    for art in articles:
        flag = (
            art['score'] >= threshold
            and not art.get('also')
            and art.get('region_slug', 'global') != 'global'
        )
        art['underreported'] = flag
        if flag:
            count += 1
    return count


def suggest_missing_feeds(articles):
    """
    Identify the 3 thinnest (region_slug × category) combinations today and
    ask Ollama to suggest RSS feeds for each gap.

    Returns feed_gaps: [{region, region_slug, category, article_count,
                          suggestions: [{name, url, reason}]}]
    Fails silently — any gap with an Ollama error gets suggestions=[].
    """
    _REGION_LABELS = {
        'ireland':      'Ireland',
        'uk':           'United Kingdom',
        'europe':       'Europe',
        'us':           'United States',
        'canada':       'Canada',
        'latam':        'Latin America',
        'asia-pacific': 'Asia-Pacific',
        'africa-me':    'Africa & Middle East',
    }

    # Count per (region_slug, category), skipping global/think-tank sources
    counts: dict = {}
    for art in articles:
        rslug = art.get('region_slug', 'global')
        if rslug == 'global':
            continue
        cat = dominant_theme(art)
        key = (rslug, cat)
        counts[key] = counts.get(key, 0) + 1

    if not counts:
        return []

    # Take the 3 thinnest (tie-break alphabetically for determinism)
    gaps = sorted(counts.items(), key=lambda x: (x[1], x[0][0], x[0][1]))[:3]

    feed_gaps = []
    for (rslug, cat), cnt in gaps:
        region_label = _REGION_LABELS.get(rslug, rslug.replace('-', ' ').title())
        prompt = (
            f"I'm building an economics news aggregator. I need more RSS feed coverage "
            f"for {cat} news from {region_label}. "
            f"Suggest 3 specific RSS feed URLs from reputable news outlets or publications "
            f"in that region covering this topic. "
            f'Return ONLY a valid JSON array with no other text: '
            f'[{{"name": "...", "url": "...", "reason": "..."}}]'
        )

        suggestions = []
        try:
            resp = requests.post(
                OLLAMA_URL,
                json={'model': OLLAMA_MODEL, 'prompt': prompt, 'stream': False},
                timeout=OLLAMA_TIMEOUT,
            )
            resp.raise_for_status()
            raw = resp.json().get('response', '').strip()
            # Strip markdown code fences Ollama sometimes adds
            raw = re.sub(r'^```(?:json)?\s*\n?', '', raw, flags=re.MULTILINE)
            raw = re.sub(r'\n?```\s*$',          '', raw, flags=re.MULTILINE)
            raw = raw.strip()
            parsed = json.loads(raw)
            if isinstance(parsed, list):
                suggestions = [
                    s for s in parsed
                    if isinstance(s, dict) and s.get('name') and s.get('url')
                ][:3]
        except Exception:
            suggestions = []

        feed_gaps.append({
            'region':        region_label,
            'region_slug':   rslug,
            'category':      cat,
            'article_count': cnt,
            'suggestions':   suggestions,
        })

    n_sugg = sum(len(g['suggestions']) for g in feed_gaps)
    print(f"  Found {len(feed_gaps)} coverage gap(s), {n_sugg} feed suggestion(s) from Ollama")
    return feed_gaps


# =============================================================================
# STEP 7 — BUILD STRUCTURED ARTICLE LIST
# =============================================================================

def build_articles(entries, keywords):
    """
    Score every entry, drop those with zero keyword matches, and return a
    sorted list of article dicts ready for summarisation and HTML output.
    """
    articles = []

    for entry in entries:
        score, themes, matched_kws = score_entry(entry, keywords)
        if score == 0:
            continue  # Not relevant to our keyword list — skip

        title  = (entry.get('title') or 'Untitled').strip()
        link   = entry.get('link', '#')
        pub_dt = entry_published_time(entry)
        source = entry.get('_source_label', 'Unknown Source')
        region = entry.get('_source_region', 'International')

        # Gather the best available body text for the AI summariser
        body = ''
        if entry.get('content'):
            body = strip_html(' '.join(c.get('value', '') for c in entry['content']))
        if not body:
            body = strip_html(entry.get('summary', '') or entry.get('description', ''))

        articles.append({
            'title':       title,
            'link':        link,
            'source':      source,
            'region':      region,
            'region_slug': REGION_SLUG_MAP.get(region, 'global'),
            'published':   pub_dt,
            'score':       score,
            'themes':      themes,
            'matched_kws': matched_kws,
            'body':        body[:2000],
            'summary':     '',
            'image':       extract_image(entry),
        })

    # Highest score first
    articles.sort(key=lambda a: a['score'], reverse=True)
    print(f"Relevant articles before deduplication: {len(articles)}")
    return articles


# =============================================================================
# STEP 8 — GENERATE AI SUMMARIES VIA OLLAMA (mistral)
# =============================================================================

def generate_summary(title, body):
    """
    Call the locally running Ollama API to produce a 2–3 sentence summary.
    If Ollama is not running or the request fails, a clear placeholder is
    returned so the rest of the digest still works fine.
    """
    if not body:
        return "No article body text was available to summarise."

    prompt = (
        "Summarise the following news article in exactly 2 to 3 concise sentences. "
        "Focus on the key economic or policy insight. "
        "Write in plain prose — no bullet points, no headings.\n\n"
        f"Title: {title}\n\n"
        f"Article text:\n{body[:1500]}"
    )

    try:
        resp = requests.post(
            OLLAMA_URL,
            json={"model": OLLAMA_MODEL, "prompt": prompt, "stream": False},
            timeout=OLLAMA_TIMEOUT
        )
        resp.raise_for_status()
        return resp.json().get('response', '').strip()

    except requests.exceptions.ConnectionError:
        # Ollama is not running — return a helpful message
        return (
            "[Ollama is not running. Start it with: ollama serve  "
            "and make sure the mistral model is pulled: ollama pull mistral]"
        )
    except Exception as exc:
        return f"[Summary unavailable: {exc}]"


def add_summaries(articles):
    """Iterate through articles and add Ollama summaries one by one."""
    total = len(articles)
    print(f"\nGenerating summaries for {total} articles via Ollama ({OLLAMA_MODEL})...")
    print("This is the slowest step — roughly 10–30 seconds per article.\n")

    for i, art in enumerate(articles, 1):
        short_title = art['title'][:70] + ('…' if len(art['title']) > 70 else '')
        print(f"  [{i:>3}/{total}] {short_title}")
        art['summary'] = generate_summary(art['title'], art['body'])

    return articles


# =============================================================================
# STEP 9 — BUILD THE HTML DIGEST
# =============================================================================

def fmt_time(dt):
    """Format a datetime for display in the digest."""
    if not dt:
        return "Time unknown"
    # strftime('%d') gives zero-padded day; lstrip('0') removes it on all platforms
    day = dt.strftime('%d').lstrip('0') or '0'
    return f"{day} {dt.strftime('%b %Y, %H:%M')} UTC"  # e.g. "6 Apr 2026, 09:15 UTC"


def dominant_theme(art):
    """
    Return the single most relevant theme for grouping this article.
    Uses THEME_PRIORITY so a story that matches both IRELAND and GLOBAL ECONOMY
    is placed under IRELAND rather than whichever comes first alphabetically.
    """
    if not art['themes']:
        return "General"
    for theme in THEME_PRIORITY:
        if theme in art['themes']:
            return theme
    return sorted(art['themes'])[0]


def slug(s):
    """Convert a string to a CSS-safe ID."""
    return re.sub(r'[^a-z0-9]+', '-', s.lower()).strip('-')


def e(s):
    """HTML-escape a value."""
    return html_module.escape(str(s))


def art_to_dict(art):
    """
    Serialise one article for the embedded JSON.
    Score is intentionally excluded — it affects ordering but is hidden from readers.
    """
    return {
        'title':       art['title'],
        'link':        art['link'],
        'source':      art['source'],
        'region':      art['region'],
        'region_slug': art.get('region_slug', 'global'),
        'published':   fmt_time(art['published']),
        'summary':     art.get('summary') or '',
        'kws':         art['matched_kws'][:5],
        'image':       art.get('image') or '',
        'also':          art.get('also', []),
        'perspectives':  art.get('perspectives', False),
        'underreported': art.get('underreported', False),
    }


def assign_subcategory(art, theme):
    """
    Return the sub-category name for an article within its theme, or None
    if none of the keyword clusters match (article will land in 'Other').
    """
    if theme not in SUBCATEGORY_RULES:
        return None
    blob = ' '.join([
        art.get('title', ''),
        art.get('summary', ''),
        ' '.join(art.get('matched_kws', [])),
    ]).lower()
    for sub_name, keywords in SUBCATEGORY_RULES[theme]:
        for kw in keywords:
            if kw in blob:
                return sub_name
    return None


def build_category_data(articles):
    """
    Group articles by dominant theme, apply sub-category rules for large groups,
    and return an ordered list of category dicts ready for JSON serialisation.
    """
    # Group by dominant theme (articles are already in score-descending order)
    theme_map = {}
    for art in articles:
        t = dominant_theme(art)
        theme_map.setdefault(t, []).append(art)

    # Build in display order, then append any extra themes alphabetically
    ordered_themes = [t for t in CATEGORY_DISPLAY_ORDER if t in theme_map]
    ordered_themes += sorted(t for t in theme_map if t not in CATEGORY_DISPLAY_ORDER)

    categories = []
    for theme in ordered_themes:
        arts = theme_map[theme]

        cat_dict = {
            'name':    theme,
            'slug':    slug(theme),
            'count':   len(arts),
            'articles': [art_to_dict(a) for a in arts],
            'subcats': [],
        }

        # Split into sub-categories only when the theme is large enough
        if len(arts) >= MIN_SUBCATS and theme in SUBCATEGORY_RULES:
            subcat_map = {}
            for art in arts:
                key = assign_subcategory(art, theme) or '__other__'
                subcat_map.setdefault(key, []).append(art)

            for sub_name, _ in SUBCATEGORY_RULES[theme]:
                sub_arts = subcat_map.get(sub_name, [])
                if sub_arts:
                    cat_dict['subcats'].append({
                        'name':     sub_name,
                        'slug':     slug(sub_name),
                        'articles': [art_to_dict(a) for a in sub_arts],
                    })

            other_arts = subcat_map.get('__other__', [])
            if other_arts:
                cat_dict['subcats'].append({
                    'name':     'Other',
                    'slug':     'other',
                    'articles': [art_to_dict(a) for a in other_arts],
                })

        categories.append(cat_dict)

    return categories



# render_card() removed — article rendering now happens in JS (see build_html)


def _count_all_keywords(articles, keywords):
    """
    Count how many articles contain each keyword (title + summary).
    Returns {keyword: count} for every keyword with count > 0.
    """
    counts = {}
    for kw_dict in keywords:
        word = kw_dict['keyword']
        count = 0
        for art in articles:
            blob = (art['title'] + ' ' + (art.get('summary') or '')).lower()
            if ' ' in word:
                if word in blob:
                    count += 1
            else:
                if re.search(r'\b' + re.escape(word) + r'\b', blob):
                    count += 1
        if count > 0:
            counts[word] = count
    return counts


def compute_top_keywords(count_dict, keywords):
    """
    Given a pre-computed {keyword: count} dict, return up to 15 entries
    sorted by count descending with at most one keyword per theme.
    """
    hits = [
        {'keyword': kd['keyword'], 'count': count_dict[kd['keyword']], 'theme': kd['theme']}
        for kd in keywords if kd['keyword'] in count_dict
    ]
    hits.sort(key=lambda x: -x['count'])
    seen_themes: set = set()
    deduped = []
    for item in hits:
        if item['theme'] not in seen_themes:
            seen_themes.add(item['theme'])
            deduped.append(item)
            if len(deduped) == 15:
                break
    return deduped


def update_keyword_history(count_dict):
    """
    Load keyword_history.json, add today's counts, prune to 30 days, save.
    Returns the updated history dict {date_str: {keyword: count}}.
    """
    today = datetime.now(timezone.utc).strftime('%Y-%m-%d')
    try:
        with open(KEYWORD_HISTORY_JSON, encoding='utf-8') as f:
            history = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        history = {}
    history[today] = count_dict
    cutoff = (datetime.now(timezone.utc) - timedelta(days=30)).strftime('%Y-%m-%d')
    history = {d: v for d, v in history.items() if d >= cutoff}
    with open(KEYWORD_HISTORY_JSON, 'w', encoding='utf-8') as f:
        json.dump(history, f, ensure_ascii=False, indent=2)
    return history


def compute_trending_keywords(count_dict, history, keywords):
    """
    For each keyword active today, compute trend_score = today / 7-day avg.
    Returns the top 8 with trend_score > 1.5, sorted by trend_score desc.
    Each entry: {keyword, todayCount, avgCount, trendScore, theme}.
    """
    today = datetime.now(timezone.utc).strftime('%Y-%m-%d')
    past_days = sorted(d for d in history if d < today)[-7:]
    kw_theme  = {kd['keyword']: kd['theme'] for kd in keywords}
    trending  = []
    for word, today_count in count_dict.items():
        if past_days:
            past_vals = [history[d].get(word, 0) for d in past_days]
            avg = sum(past_vals) / len(past_days)
        else:
            avg = 1.0
        trend_score = today_count / max(avg, 1.0)
        if trend_score > 1.5:
            trending.append({
                'keyword':    word,
                'todayCount': today_count,
                'avgCount':   round(avg, 1),
                'trendScore': round(trend_score, 2),
                'theme':      kw_theme.get(word, 'General'),
            })
    trending.sort(key=lambda x: -x['trendScore'])
    return trending


def _find_prev_digest():
    """Return the filename (e.g. 'digest-2026-05-03.html') of the most recent
    archived digest that predates today, or None if none exists."""
    archive_dir = os.path.join(BASE_DIR, 'archive')
    if not os.path.isdir(archive_dir):
        return None
    today_file = 'digest-' + datetime.now(timezone.utc).strftime('%Y-%m-%d') + '.html'
    pat = re.compile(r'^digest-\d{4}-\d{2}-\d{2}\.html$')
    prior = sorted(f for f in os.listdir(archive_dir) if pat.match(f) and f < today_file)
    return prior[-1] if prior else None


def build_archive_index(archive_dir):
    """Return a complete HTML string for archive/index.html."""
    pat = re.compile(r'^digest-(\d{4}-\d{2}-\d{2})\.html$')
    files = sorted(
        (f for f in os.listdir(archive_dir) if pat.match(f)),
        reverse=True,
    )
    items = []
    for fname in files:
        date_str = pat.match(fname).group(1)
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        day      = str(date_obj.day)
        label    = date_obj.strftime(f'%A, {day} %B %Y')
        count_str = ''
        try:
            with open(os.path.join(archive_dir, fname), encoding='utf-8') as fh:
                tail = fh.read()[-3000:]
            m = re.search(r'(\d+) articles from the last', tail)
            if m:
                count_str = f' &middot; {m.group(1)} articles'
        except Exception:
            pass
        items.append(f'    <li><a href="{fname}">{label}{count_str}</a></li>')
    n     = len(files)
    rows  = '\n'.join(items) if items else '    <li><em>No digests archived yet.</em></li>'
    plural = '' if n == 1 else 's'
    return f'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Digest Archive</title>
  <style>
    *, *::before, *::after {{ box-sizing: border-box; }}
    body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
           background: #0f172a; color: #e2e8f0; padding: 2rem 1.5rem;
           max-width: 700px; margin: 0 auto; }}
    a {{ color: inherit; text-decoration: none; }}
    .back {{ display: inline-block; margin-bottom: 1.5rem; color: #94a3b8;
             font-size: 0.85rem; }}
    .back:hover {{ color: #f87171; }}
    h1 {{ color: #f87171; margin: 0 0 0.2rem; font-size: 1.4rem; }}
    .sub {{ color: #94a3b8; font-size: 0.85rem; margin: 0 0 1.75rem; }}
    ul {{ list-style: none; padding: 0; margin: 0; }}
    li {{ border-bottom: 1px solid #1e293b; }}
    li a {{ display: block; padding: 0.65rem 0; }}
    li a:hover {{ color: #f87171; }}
  </style>
</head>
<body>
  <a href="../digest.html" class="back">&#8592; Back to today&rsquo;s digest</a>
  <h1>&#128193; Digest Archive</h1>
  <p class="sub">{n} digest{plural} archived</p>
  <ul>
{rows}
  </ul>
</body>
</html>'''


def build_trends_page(history, trending_keywords):
    """Generate archive/trends.html — a standalone keyword-trends page."""
    archive_dir = os.path.join(BASE_DIR, 'archive')
    os.makedirs(archive_dir, exist_ok=True)

    today   = datetime.now(timezone.utc).strftime('%Y-%m-%d')
    today_obj = datetime.strptime(today, '%Y-%m-%d')
    day_str   = str(today_obj.day)
    now_str   = today_obj.strftime(f'%A, {day_str} %B %Y')

    # Last 7 calendar days (excluding today)
    past_7 = [(today_obj - timedelta(days=i)).strftime('%Y-%m-%d') for i in range(1, 8)]
    past_7.reverse()   # oldest → newest

    # ── Trending rows ────────────────────────────────────────────────────────
    trend_rows_html = ''
    if trending_keywords:
        max_score = max(k['trendScore'] for k in trending_keywords)
        for kw in trending_keywords:
            bar_w   = max(4, round((kw['trendScore'] / max_score) * 120))
            pct     = round((kw['trendScore'] - 1) * 100)
            sign    = '+' if pct >= 0 else ''
            trend_rows_html += f'''
      <tr>
        <td class="kw-name">{html_module.escape(kw["keyword"])}</td>
        <td class="num">{kw["todayCount"]}</td>
        <td class="num">{kw["avgCount"]}</td>
        <td class="num score">{kw["trendScore"]:.2f}×</td>
        <td class="bar-cell">
          <div class="spark-track"><div class="spark-bar" style="width:{bar_w}px"></div></div>
        </td>
        <td class="pct">{sign}{pct}%</td>
        <td class="theme-tag">{html_module.escape(kw["theme"])}</td>
      </tr>'''
    else:
        trend_rows_html = '<tr><td colspan="7" class="empty">No spiking keywords today — need at least 2 days of history.</td></tr>'

    # ── Recent history table ─────────────────────────────────────────────────
    # Collect all keywords that appeared on any of the last 7 days
    all_kws: set = set()
    for d in past_7:
        all_kws.update(history.get(d, {}).keys())
    # Also include today's trending keywords
    for kw in trending_keywords:
        all_kws.add(kw['keyword'])

    today_counts = {kw['keyword']: kw['todayCount'] for kw in trending_keywords}

    # Sort by today's count desc, then alpha
    sorted_kws = sorted(all_kws, key=lambda w: (-today_counts.get(w, 0), w))

    # Day column headers (short format)
    day_headers = ''
    for d in past_7:
        d_obj = datetime.strptime(d, '%Y-%m-%d')
        day_headers += f'<th>{d_obj.strftime("%-d %b") if os.name != "nt" else d_obj.strftime("%#d %b")}</th>'
    day_headers += '<th>Today</th>'

    history_rows_html = ''
    for word in sorted_kws:
        cells = ''
        for d in past_7:
            cnt = history.get(d, {}).get(word, 0)
            cells += f'<td class="num{"" if cnt else " zero"}">{cnt if cnt else ""}</td>'
        today_cnt = today_counts.get(word, history.get(today, {}).get(word, 0))
        cells += f'<td class="num today-col">{today_cnt if today_cnt else ""}</td>'
        history_rows_html += f'<tr><td class="kw-name">{html_module.escape(word)}</td>{cells}</tr>'

    if not history_rows_html:
        history_rows_html = '<tr><td colspan="9" class="empty">No keyword history yet.</td></tr>'

    return f'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Keyword Trends — {now_str}</title>
  <style>
    *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
      background: #0f172a; color: #e2e8f0;
      padding: 2rem 1.5rem; max-width: 900px; margin: 0 auto;
    }}
    a {{ color: inherit; text-decoration: none; }}
    .back {{ display: inline-block; margin-bottom: 1.5rem; color: #94a3b8; font-size: 0.85rem; }}
    .back:hover {{ color: #f87171; }}
    h1 {{ color: #f87171; margin: 0 0 0.2rem; font-size: 1.4rem; }}
    h2 {{ color: #f87171; font-size: 1rem; margin: 2rem 0 0.75rem; }}
    .sub {{ color: #94a3b8; font-size: 0.85rem; margin: 0 0 2rem; }}
    table {{ width: 100%; border-collapse: collapse; font-size: 0.83rem; }}
    th {{
      text-align: left; color: #94a3b8; font-size: 0.72rem; font-weight: 600;
      text-transform: uppercase; letter-spacing: 0.06em;
      border-bottom: 1px solid #1e293b; padding: 0.4rem 0.5rem;
    }}
    td {{ padding: 0.4rem 0.5rem; border-bottom: 1px solid #1e293b; vertical-align: middle; }}
    tr:hover td {{ background: #1e293b; }}
    .kw-name {{ font-weight: 500; }}
    .num {{ text-align: right; font-variant-numeric: tabular-nums; color: #cbd5e1; }}
    .zero {{ color: #334155; }}
    .today-col {{ color: #f87171; font-weight: 700; }}
    .score {{ color: #f87171; font-weight: 700; }}
    .pct {{ color: #4ade80; font-weight: 700; text-align: right; }}
    .theme-tag {{ color: #94a3b8; font-size: 0.72rem; }}
    .bar-cell {{ padding: 0.4rem 0.5rem; }}
    .spark-track {{
      width: 120px; height: 6px; background: #1e293b;
      border-radius: 3px; overflow: hidden;
    }}
    .spark-bar {{ height: 100%; background: #f87171; border-radius: 3px; }}
    .empty {{ color: #64748b; font-style: italic; padding: 1rem 0.5rem; }}
  </style>
</head>
<body>
  <a href="../digest.html" class="back">&#8592; Back to today&rsquo;s digest</a>
  <h1>&#128200; Keyword Trends</h1>
  <p class="sub">Generated {now_str}</p>

  <h2>Spiking Today (trend score &gt; 1.5×)</h2>
  <table>
    <thead>
      <tr>
        <th>Keyword</th>
        <th style="text-align:right">Today</th>
        <th style="text-align:right">7-day avg</th>
        <th style="text-align:right">Score</th>
        <th>Sparkline</th>
        <th style="text-align:right">vs avg</th>
        <th>Theme</th>
      </tr>
    </thead>
    <tbody>{trend_rows_html}
    </tbody>
  </table>

  <h2>Recent History (last 7 days)</h2>
  <table>
    <thead>
      <tr>
        <th>Keyword</th>
        {day_headers}
      </tr>
    </thead>
    <tbody>{history_rows_html}
    </tbody>
  </table>
</body>
</html>'''


def build_feed_dashboard(sources, feed_article_counts, disabled_count):
    """Generate archive/feeds.html — per-feed contribution dashboard."""
    today_obj = datetime.now(timezone.utc)
    day_str   = str(today_obj.day)
    now_str   = today_obj.strftime(f'%A, {day_str} %B %Y')

    # Enrich each source with its article count for today
    rows = []
    for src in sources:
        count = feed_article_counts.get(src['label'], 0)
        rows.append({
            'label':  src['label'],
            'region': src['region'],
            'count':  count,
        })
    rows.sort(key=lambda r: (-r['count'], r['label'].lower()))

    total_feeds  = len(rows)
    contributing = sum(1 for r in rows if r['count'] > 0)
    silent       = total_feeds - contributing

    # Top 10 by count, for bar chart
    top10    = [r for r in rows if r['count'] > 0][:10]
    max_cnt  = max((r['count'] for r in top10), default=1)

    bars_html = ''
    for r in top10:
        pct = round(r['count'] / max_cnt * 100)
        bars_html += f'''
    <div class="bar-row">
      <span class="bar-label">{html_module.escape(r["label"])}</span>
      <div class="bar-track"><div class="bar-fill" style="width:{pct}%"></div></div>
      <span class="bar-count">{r["count"]}</span>
    </div>'''
    if not bars_html:
        bars_html = '\n    <p class="empty">No articles today.</p>'

    # Main table rows
    table_rows = ''
    for r in rows:
        if r['count'] > 0:
            status = '<span class="badge-active">Active</span>'
            cnt    = str(r['count'])
        else:
            status = '<span class="badge-silent">Silent</span>'
            cnt    = '&mdash;'
        table_rows += f'''
      <tr>
        <td class="feed-name">{html_module.escape(r["label"])}</td>
        <td class="feed-region">{html_module.escape(r["region"])}</td>
        <td class="feed-count">{cnt}</td>
        <td>{status}</td>
      </tr>'''

    return f'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Feed Dashboard — {now_str}</title>
  <style>
    *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
      background: #0f172a; color: #e2e8f0;
      padding: 2rem 1.5rem; max-width: 900px; margin: 0 auto;
    }}
    a {{ color: inherit; text-decoration: none; }}
    .back {{ display: inline-block; margin-bottom: 1.5rem; color: #94a3b8; font-size: 0.85rem; }}
    .back:hover {{ color: #f87171; }}
    h1 {{ color: #f87171; margin: 0 0 0.2rem; font-size: 1.4rem; }}
    h2 {{ color: #f87171; font-size: 1rem; margin: 2rem 0 0.75rem; border-bottom: 1px solid #1e293b; padding-bottom: 0.4rem; }}
    .sub {{ color: #94a3b8; font-size: 0.85rem; margin: 0 0 2rem; }}
    /* Stats row */
    .stats-row {{
      display: flex; flex-wrap: wrap; gap: 1rem; margin-bottom: 2rem;
    }}
    .stat-card {{
      background: #1e293b; border-radius: 8px; padding: 0.75rem 1.25rem; flex: 1; min-width: 120px;
    }}
    .stat-num {{ font-size: 1.6rem; font-weight: 800; color: #f87171; line-height: 1; }}
    .stat-label {{ font-size: 0.72rem; color: #94a3b8; text-transform: uppercase; letter-spacing: 0.06em; margin-top: 0.2rem; }}
    /* Bar chart */
    .bar-row {{
      display: flex; align-items: center; gap: 0.6rem;
      margin-bottom: 0.45rem; font-size: 0.82rem;
    }}
    .bar-label {{ width: 180px; flex-shrink: 0; color: #cbd5e1; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }}
    .bar-track {{ flex: 1; background: #1e293b; border-radius: 3px; height: 10px; overflow: hidden; }}
    .bar-fill {{ height: 100%; background: #f87171; border-radius: 3px; }}
    .bar-count {{ width: 2.5rem; text-align: right; color: #f87171; font-weight: 700; font-size: 0.82rem; }}
    /* Main table */
    table {{ width: 100%; border-collapse: collapse; font-size: 0.83rem; }}
    th {{
      text-align: left; color: #94a3b8; font-size: 0.72rem; font-weight: 600;
      text-transform: uppercase; letter-spacing: 0.06em;
      border-bottom: 1px solid #1e293b; padding: 0.4rem 0.5rem;
    }}
    td {{ padding: 0.45rem 0.5rem; border-bottom: 1px solid #1e293b; vertical-align: middle; }}
    tr:hover td {{ background: #1e293b; }}
    .feed-name {{ font-weight: 500; }}
    .feed-region {{ color: #94a3b8; font-size: 0.78rem; }}
    .feed-count {{ text-align: right; font-variant-numeric: tabular-nums; color: #cbd5e1; }}
    .badge-active {{
      display: inline-block; font-size: 0.68rem; font-weight: 700;
      padding: 1px 7px; border-radius: 99px;
      background: #dcfce7; color: #166534;
    }}
    .badge-silent {{
      display: inline-block; font-size: 0.68rem; font-weight: 700;
      padding: 1px 7px; border-radius: 99px;
      background: #fef3c7; color: #92400e;
    }}
    .empty {{ color: #64748b; font-style: italic; font-size: 0.85rem; }}
  </style>
</head>
<body>
  <a href="../digest.html" class="back">&#8592; Back to today&rsquo;s digest</a>
  <h1>&#128225; Feed Dashboard</h1>
  <p class="sub">Generated {now_str}</p>

  <div class="stats-row">
    <div class="stat-card"><div class="stat-num">{total_feeds}</div><div class="stat-label">Active feeds</div></div>
    <div class="stat-card"><div class="stat-num">{contributing}</div><div class="stat-label">Contributing today</div></div>
    <div class="stat-card"><div class="stat-num">{silent}</div><div class="stat-label">Silent today</div></div>
    <div class="stat-card"><div class="stat-num">{disabled_count}</div><div class="stat-label">Disabled (##)</div></div>
  </div>

  <h2>Top Performers Today</h2>
  <div class="bar-chart">{bars_html}
  </div>

  <h2>All Feeds</h2>
  <table>
    <thead>
      <tr>
        <th>Feed</th>
        <th>Region</th>
        <th style="text-align:right">Articles today</th>
        <th>Status</th>
      </tr>
    </thead>
    <tbody>{table_rows}
    </tbody>
  </table>
</body>
</html>'''


def archive_digest(digest_html, total):
    """Copy today's digest into archive/ and regenerate the archive index."""
    archive_dir = os.path.join(BASE_DIR, 'archive')
    os.makedirs(archive_dir, exist_ok=True)
    today_str    = datetime.now(timezone.utc).strftime('%Y-%m-%d')
    archive_path = os.path.join(archive_dir, f'digest-{today_str}.html')
    with open(archive_path, 'w', encoding='utf-8') as f:
        f.write(digest_html)
    print(f"  Archived → {archive_path}")
    index_html = build_archive_index(archive_dir)
    index_path = os.path.join(archive_dir, 'index.html')
    with open(index_path, 'w', encoding='utf-8') as f:
        f.write(index_html)
    print(f"  Archive index → {index_path}")


def build_html(articles, prev_digest=None, top_keywords=None,
               feeds_checked=0, trending_keywords=None, feed_gaps=None):
    """Assemble the complete HTML digest from all articles."""

    _now     = datetime.now(timezone.utc)
    day_str  = _now.strftime('%d').lstrip('0') or '0'
    now_str  = _now.strftime(f"%A, {day_str} %B %Y at %H:%M GMT")
    total    = len(articles)

    # Build structured category/sub-category data
    categories = build_category_data(articles)

    # Top N articles (score-sorted; score intentionally absent from output dict)
    top_arts = [art_to_dict(a) for a in articles[:TOP_N]]

    # Prev/next nav — nextDigest is always null on a fresh generation
    prev_btn_attrs = (
        f"onclick=\"window.location.href='archive/{prev_digest}'\""
        if prev_digest else 'disabled'
    )

    # Embed all data as JSON — JS does the rendering
    data_json = json.dumps({
        'generated':     now_str,
        'totalArticles': total,
        'hours':         HOURS_BACK,
        'topStories':    top_arts,
        'categories':    categories,
        'prevDigest':    f'archive/{prev_digest}' if prev_digest else None,
        'nextDigest':    None,
        'topKeywords':       top_keywords or [],
        'trendingKeywords':  trending_keywords or [],
        'feedGaps':          feed_gaps or [],
        'stats': {
            'totalArticles':     total,
            'feedsChecked':      feeds_checked,
            'hoursBack':         HOURS_BACK,
            'model':             OLLAMA_MODEL,
            'underreportedCount': sum(
                1 for a in articles if a.get('underreported')
            ),
        },
    }, ensure_ascii=False)

    # Build nav tabs (Python-side so they're in the DOM immediately, no flash)
    nav_tabs_html = '<button class="nav-tab active" data-cat="__home__">All Topics</button>\n'
    nav_tabs_html += '<button class="nav-tab starred" data-cat="__top__">&#9733; Top Stories</button>\n'
    nav_tabs_html += '<button class="nav-tab starred" data-cat="__saved__" id="saved-tab">&#9733; Saved</button>\n'
    for cat in categories:
        nav_tabs_html += (
            f'<button class="nav-tab" data-cat="{e(cat["slug"])}">'
            f'{e(cat["name"])} <span class="tab-count">{cat["count"]}</span>'
            f'</button>\n'
        )


    # Extra variables for the new editorial template
    live_time_str = _now.strftime('%H:%M')
    hours_back    = HOURS_BACK
    model_str     = OLLAMA_MODEL

    # ── Full HTML document ────────────────────────────────────────────────────
    return f"""
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Economic Digest — {now_str}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Source+Serif+4:opsz,wght@8..60,400;8..60,500;8..60,600;8..60,700;8..60,800&family=Inter:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>
  /* ─── Reset ────────────────────────────────────────────────────── */
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}

  /* ─── Tokens ───────────────────────────────────────────────────── */
  :root {{
    --paper:        #ffffff;
    --canvas:       #f4f1ec;          /* warm off-white page bg  */
    --rule:         #1a1a1a;          /* hairline rule colour    */
    --rule-soft:    #d8d4cc;
    --ink:          #111111;
    --ink-2:        #2a2a2a;
    --muted:        #6e6a63;
    --muted-2:      #8a857c;
    --accent:       #8a2a2a;          /* deep oxblood accent     */
    --tag-bg:       #efece6;
    --tag-ink:      #2a2a2a;
    --serif:        "Source Serif 4", "Georgia", "Times New Roman", serif;
    --sans:         "Inter", -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    --mono:         "JetBrains Mono", ui-monospace, Menlo, monospace;
    --max:          1280px;
    --gap:          24px;
  }}
  body.dark {{
    --paper:        #161616;
    --canvas:       #0e0e0e;
    --rule:         #f1f1f1;
    --rule-soft:    #2c2c2c;
    --ink:          #f1f1f1;
    --ink-2:        #d8d6d0;
    --muted:        #9a958c;
    --muted-2:      #7a766d;
    --accent:       #d98b8b;
    --tag-bg:       #1f1f1f;
    --tag-ink:      #d8d6d0;
  }}

  html, body {{ background: var(--canvas); color: var(--ink); }}
  body {{
    font-family: var(--sans);
    font-size: 15px;
    line-height: 1.5;
    -webkit-font-smoothing: antialiased;
  }}
  a {{ color: inherit; text-decoration: none; }}
  a:hover {{ text-decoration: underline; text-underline-offset: 2px; text-decoration-thickness: 1px; }}

  /* ─── Top bar ─────────────────────────────────────────────────── */
  .topbar {{
    background: var(--paper);
    border-bottom: 1px solid var(--rule-soft);
  }}
  .topbar-inner {{
    max-width: var(--max);
    margin: 0 auto;
    padding: 14px 24px;
    display: grid;
    grid-template-columns: 1fr auto 1fr;
    align-items: center;
    gap: 16px;
  }}
  .topbar-left, .topbar-right {{
    display: flex;
    align-items: center;
    gap: 10px;
  }}
  .topbar-right {{ justify-content: flex-end; }}
  .icon-btn {{
    width: 34px; height: 34px;
    display: inline-flex; align-items: center; justify-content: center;
    background: transparent;
    border: 1px solid transparent;
    border-radius: 4px;
    color: var(--ink);
    cursor: pointer;
    font-size: 16px;
  }}
  .icon-btn:hover {{ background: var(--tag-bg); }}
  .live-pill {{
    display: inline-flex; align-items: center; gap: 6px;
    padding: 5px 10px;
    border: 1px solid var(--rule-soft);
    border-radius: 999px;
    font-family: var(--sans);
    font-size: 11px;
    font-weight: 600;
    letter-spacing: 0.04em;
    color: var(--ink-2);
    cursor: default;
  }}
  .live-dot {{
    width: 7px; height: 7px;
    border-radius: 50%;
    background: var(--accent);
  }}
  .wordmark {{
    text-align: center;
    font-family: var(--serif);
    font-weight: 800;
    font-size: 26px;
    letter-spacing: 0.005em;
    line-height: 1;
    color: var(--ink);
    user-select: none;
  }}
  .wordmark .ed {{
    display: inline-block;
    border-top: 2px solid var(--ink);
    border-bottom: 2px solid var(--ink);
    padding: 4px 8px;
    margin-right: 2px;
  }}
  .wordmark .nd {{
    font-style: italic;
    font-weight: 600;
    margin-left: 2px;
  }}
  .pill-btn {{
    padding: 7px 14px;
    border-radius: 999px;
    font-family: var(--sans);
    font-size: 12px;
    font-weight: 600;
    cursor: pointer;
    border: 1px solid var(--ink);
    background: var(--ink);
    color: var(--paper);
  }}
  .pill-btn.ghost {{
    background: transparent;
    color: var(--ink);
    border: 1px solid var(--rule-soft);
  }}
  .pill-btn.ghost:hover {{ border-color: var(--ink); }}

  /* ─── Search expand ─────────────────────────────────────────── */
  .search-wrap {{
    display: flex; align-items: center;
  }}
  .search-wrap input {{
    width: 0;
    padding: 6px 0;
    border: 0;
    border-bottom: 1px solid transparent;
    background: transparent;
    color: var(--ink);
    font-family: var(--sans);
    font-size: 13px;
    outline: none;
    transition: width .25s ease, padding .25s ease, border-color .25s ease;
  }}
  .search-wrap.open input {{
    width: 200px;
    padding: 6px 8px;
    border-bottom-color: var(--ink);
  }}

  /* ─── Primary nav ─────────────────────────────────────────────── */
  .navbar {{
    position: sticky;
    top: 0;
    z-index: 100;
    background: var(--paper);
    border-bottom: 1px solid var(--rule);
  }}
  .nav-inner {{
    max-width: var(--max);
    margin: 0 auto;
    padding: 0 24px;
    display: flex;
    align-items: stretch;
    overflow-x: auto;
    scrollbar-width: none;
    gap: 0;
  }}
  .nav-inner::-webkit-scrollbar {{ display: none; }}
  .nav-tab {{
    flex-shrink: 0;
    background: transparent;
    border: 0;
    padding: 14px 14px;
    font-family: var(--sans);
    font-size: 11px;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: var(--ink-2);
    cursor: pointer;
    position: relative;
    white-space: nowrap;
  }}
  .nav-tab:hover {{ color: var(--ink); }}
  .nav-tab.active {{ color: var(--ink); }}
  .nav-tab.active::after {{
    content: "";
    position: absolute;
    left: 14px; right: 14px;
    bottom: -1px;
    height: 3px;
    background: var(--ink);
  }}
  .nav-tab .tab-count {{
    margin-left: 5px;
    font-size: 9px;
    font-weight: 600;
    color: var(--muted);
    letter-spacing: 0;
  }}
  .nav-tab.starred {{ color: var(--accent); }}
  .nav-tab.starred.active {{ color: var(--accent); }}
  .nav-tab.starred.active::after {{ background: var(--accent); }}

  /* ─── Region strip ──────────────────────────────────────────── */
  .regionbar {{
    background: var(--canvas);
    border-bottom: 1px solid var(--rule-soft);
  }}
  .region-inner {{
    max-width: var(--max);
    margin: 0 auto;
    padding: 10px 24px;
    display: flex;
    align-items: center;
    gap: 6px;
    overflow-x: auto;
    scrollbar-width: none;
  }}
  .region-inner::-webkit-scrollbar {{ display: none; }}
  .region-label {{
    font-family: var(--sans);
    font-size: 10px;
    text-transform: uppercase;
    letter-spacing: 0.12em;
    color: var(--muted);
    margin-right: 6px;
    white-space: nowrap;
  }}
  .region-tab {{
    background: transparent;
    border: 1px solid var(--rule-soft);
    border-radius: 999px;
    padding: 4px 12px;
    font-family: var(--sans);
    font-size: 12px;
    color: var(--ink-2);
    cursor: pointer;
    white-space: nowrap;
  }}
  .region-tab:hover {{ border-color: var(--ink); }}
  .region-tab.active {{
    background: var(--ink);
    color: var(--paper);
    border-color: var(--ink);
  }}

  /* ─── Page wrapper ────────────────────────────────────────── */
  main {{
    max-width: var(--max);
    margin: 0 auto;
    padding: 32px 24px 80px;
    background: transparent;
  }}

  /* ─── Page title (section header) ─────────────────────────── */
  .page-head {{
    margin-bottom: 28px;
  }}
  .page-title {{
    font-family: var(--serif);
    font-weight: 700;
    font-size: 56px;
    line-height: 1.05;
    letter-spacing: -0.015em;
    color: var(--ink);
  }}
  .page-title-sub {{
    margin-top: 6px;
    font-family: var(--sans);
    font-size: 13px;
    color: var(--muted);
  }}
  .page-title-count {{
    font-family: var(--serif);
    font-weight: 400;
    font-style: italic;
    font-size: 28px;
    color: var(--muted);
    margin-left: 12px;
  }}
  .back-btn {{
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: transparent;
    border: 0;
    padding: 0;
    margin-bottom: 14px;
    font-family: var(--sans);
    font-size: 12px;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    color: var(--muted);
    cursor: pointer;
  }}
  .back-btn:hover {{ color: var(--ink); }}

  /* ─── Themes bar ─────────────────────────────────────────── */
  .themes-bar {{
    margin: 0 0 32px;
    padding: 14px 0;
    border-top: 1px solid var(--rule-soft);
    border-bottom: 1px solid var(--rule-soft);
    display: flex;
    align-items: center;
    justify-content: space-between;
    gap: 24px;
    flex-wrap: wrap;
  }}
  .themes-left {{ display: flex; align-items: center; gap: 12px; flex-wrap: wrap; }}
  .themes-label {{
    font-family: var(--sans);
    font-size: 10px;
    text-transform: uppercase;
    letter-spacing: 0.14em;
    color: var(--muted);
    font-weight: 700;
  }}
  .themes-pills {{ display: flex; gap: 6px; flex-wrap: wrap; }}
  .theme-pill {{
    background: transparent;
    border: 1px solid var(--rule-soft);
    border-radius: 999px;
    padding: 3px 10px;
    font-family: var(--sans);
    font-size: 12px;
    color: var(--ink-2);
    cursor: pointer;
  }}
  .theme-pill:hover {{ border-color: var(--ink); }}
  .themes-right {{
    font-family: var(--mono);
    font-size: 11px;
    color: var(--muted);
    letter-spacing: 0.02em;
  }}

  /* ─── Hero (top stories layout) ───────────────────────────── */
  .hero {{
    display: grid;
    grid-template-columns: 1fr 1.3fr 1fr;
    gap: var(--gap);
    padding-bottom: 32px;
    border-bottom: 1px solid var(--rule-soft);
    margin-bottom: 32px;
  }}
  .hero-flank, .hero-center {{ display: flex; flex-direction: column; }}
  .hero-flank .story-card {{ padding-top: 8px; }}
  .hero-center .story-card .thumb {{ aspect-ratio: 16 / 11; }}
  .hero-center .headline {{ font-size: 28px; }}
  .hero-flank .headline {{ font-size: 22px; }}
  .hero-flank .thumb {{ display: none; }}

  @media (max-width: 900px) {{
    .hero {{ grid-template-columns: 1fr; }}
    .hero-flank .thumb {{ display: block; }}
    .hero-center .headline, .hero-flank .headline {{ font-size: 24px; }}
  }}

  /* ─── Story card (BBC-ish, original treatment) ─────────────── */
  .story-card {{
    display: flex;
    flex-direction: column;
    gap: 10px;
    background: transparent;
    border: 0;
    border-radius: 0;
    box-shadow: none;
    position: relative;
  }}
  .thumb-wrap {{ position: relative; overflow: hidden; }}
  .thumb {{
    width: 100%;
    aspect-ratio: 16 / 9;
    background: linear-gradient(135deg, #efece6 0%, #e3dfd6 100%);
    background-size: cover;
    background-position: center;
    display: block;
  }}
  body.dark .thumb {{
    background: linear-gradient(135deg, #2a2a2a 0%, #1f1f1f 100%);
  }}
  .thumb.no-image::after {{
    content: attr(data-label);
    position: absolute;
    inset: 0;
    display: flex;
    align-items: center;
    justify-content: center;
    font-family: var(--mono);
    font-size: 11px;
    color: var(--muted);
    letter-spacing: 0.1em;
    text-transform: uppercase;
  }}
  .bookmark-btn {{
    position: absolute;
    top: 8px; right: 8px;
    width: 28px; height: 28px;
    display: inline-flex; align-items: center; justify-content: center;
    background: rgba(255,255,255,0.92);
    color: var(--ink);
    border: 0;
    border-radius: 0;
    font-size: 14px;
    cursor: pointer;
    line-height: 1;
    transition: background .15s;
  }}
  body.dark .bookmark-btn {{ background: rgba(22,22,22,0.92); color: var(--ink); }}
  .bookmark-btn:hover {{ background: var(--paper); }}
  body.dark .bookmark-btn:hover {{ background: var(--paper); }}
  .bookmark-btn.saved {{ color: var(--accent); }}

  .headline {{
    font-family: var(--serif);
    font-weight: 700;
    font-size: 18px;
    line-height: 1.18;
    letter-spacing: -0.01em;
    color: var(--ink);
    text-wrap: balance;
  }}
  .headline a:hover {{ text-decoration: underline; text-decoration-thickness: 1px; text-underline-offset: 3px; }}
  .teaser {{
    font-family: var(--serif);
    font-weight: 400;
    font-size: 14.5px;
    line-height: 1.45;
    color: var(--ink-2);
    margin-top: -2px;
  }}
  .teaser.collapsed {{ display: none; }}
  .read-summary {{
    font-family: var(--sans);
    font-size: 11px;
    font-weight: 600;
    letter-spacing: 0.06em;
    text-transform: uppercase;
    color: var(--muted);
    background: transparent;
    border: 0;
    padding: 0;
    cursor: pointer;
    align-self: flex-start;
  }}
  .read-summary:hover {{ color: var(--ink); }}
  .meta-row {{
    display: flex;
    align-items: center;
    gap: 8px;
    font-family: var(--sans);
    font-size: 12px;
    color: var(--muted);
    flex-wrap: wrap;
  }}
  .meta-region {{
    font-family: var(--sans);
    font-size: 10px;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    color: var(--muted);
  }}
  .meta-sep {{ color: var(--muted-2); }}
  .meta-source {{ color: var(--ink-2); font-weight: 500; }}
  .top-flag {{
    display: inline-block;
    font-family: var(--sans);
    font-size: 10px;
    font-weight: 700;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: var(--accent);
    margin-bottom: 2px;
  }}
  .solo-flag {{
    font-family: var(--sans);
    font-size: 10px;
    font-weight: 600;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    color: var(--muted);
    border: 1px solid var(--rule-soft);
    padding: 1px 6px;
    border-radius: 2px;
  }}

  /* ─── Also-covering links ────────────────────────────────── */
  .also-link {{
    margin-top: 4px;
    font-family: var(--sans);
    font-size: 12px;
    color: var(--muted);
    cursor: pointer;
    background: transparent;
    border: 0;
    padding: 0;
    text-align: left;
  }}
  .also-link:hover {{ color: var(--ink); }}
  .also-list {{
    margin-top: 6px;
    padding: 10px 12px;
    border-left: 2px solid var(--rule-soft);
    list-style: none;
    font-family: var(--sans);
    font-size: 12.5px;
  }}
  .also-list li {{ padding: 3px 0; }}
  .also-list a {{ color: var(--ink-2); }}
  .also-list .also-src {{ color: var(--muted); font-size: 11px; }}
  .persp-row {{
    display: flex; align-items: baseline; gap: 8px;
    padding: 4px 0;
    font-family: var(--sans);
    font-size: 12.5px;
    flex-wrap: wrap;
  }}
  .persp-row .meta-region {{ min-width: 80px; }}
  .persp-snippet {{ color: var(--muted); }}

  /* ─── Section block (category preview) ─────────────────────── */
  .home-grid {{ display: flex; flex-direction: column; gap: 48px; }}
  .cat-section {{ padding-top: 0; }}
  .cat-rule {{
    border: 0;
    border-top: 1px solid var(--rule);
    margin: 0 0 16px 0;
  }}
  .cat-section-head {{
    display: flex;
    align-items: baseline;
    justify-content: space-between;
    gap: 16px;
    margin-bottom: 18px;
  }}
  .cat-section-title {{
    font-family: var(--sans);
    font-size: 11px;
    font-weight: 700;
    letter-spacing: 0.12em;
    text-transform: uppercase;
    color: var(--ink);
  }}
  .cat-section-title .cnt {{
    margin-left: 8px;
    font-weight: 500;
    color: var(--muted);
    letter-spacing: 0.04em;
  }}
  .see-all-btn {{
    background: transparent;
    border: 0;
    padding: 0;
    font-family: var(--sans);
    font-size: 12px;
    font-weight: 600;
    color: var(--ink-2);
    cursor: pointer;
  }}
  .see-all-btn:hover {{ color: var(--accent); }}

  /* ─── Story grids ─────────────────────────────────────────── */
  .story-grid {{
    display: grid;
    grid-template-columns: repeat(4, 1fr);
    gap: var(--gap) calc(var(--gap) * 1);
    list-style: none;
    margin: 0;
    padding: 0;
  }}
  /* On home grid we show 4 across; on category page we show 3 across with summary */
  .story-grid.three {{ grid-template-columns: repeat(3, 1fr); }}
  .story-grid.one   {{ grid-template-columns: 1fr; }}
  @media (max-width: 1024px) {{
    .story-grid {{ grid-template-columns: repeat(2, 1fr); }}
    .story-grid.three {{ grid-template-columns: repeat(2, 1fr); }}
  }}
  @media (max-width: 600px) {{
    .story-grid, .story-grid.three {{ grid-template-columns: 1fr; }}
  }}

  .story-card.search-hidden {{ display: none; }}
  .cat-section.search-hidden {{ display: none; }}
  .subcat-section.search-hidden {{ display: none; }}

  .subcat-section {{ margin-bottom: 36px; }}
  .subcat-heading {{
    font-family: var(--sans);
    font-size: 11px;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: var(--ink-2);
    margin-bottom: 14px;
    padding-bottom: 8px;
    border-bottom: 1px solid var(--rule-soft);
  }}
  .subcat-heading .subcat-count {{ color: var(--muted); margin-left: 6px; font-weight: 500; }}

  /* No results */
  #no-results {{
    display: none;
    max-width: var(--max);
    margin: 60px auto;
    padding: 0 24px;
    font-family: var(--serif);
    font-size: 22px;
    text-align: center;
    color: var(--muted);
  }}

  /* ─── Footer / gaps ───────────────────────────────────────── */
  footer {{
    background: #111;
    color: #d4d2cc;
    padding: 48px 24px 72px;
    margin-top: 60px;
  }}
  body.dark footer {{ background: #050505; border-top: 1px solid #1f1f1f; }}
  .footer-inner {{
    max-width: var(--max);
    margin: 0 auto;
    display: grid;
    grid-template-columns: 1fr auto;
    gap: 24px;
    align-items: end;
  }}
  .footer-brand {{
    font-family: var(--serif);
    font-size: 22px;
    font-weight: 700;
    color: #fff;
    letter-spacing: 0.005em;
    margin-bottom: 8px;
  }}
  .footer-meta {{
    font-family: var(--sans);
    font-size: 12px;
    color: #8d8a83;
    line-height: 1.6;
  }}
  .gaps-toggle {{
    background: transparent;
    color: #d4d2cc;
    border: 1px solid #3a3a3a;
    border-radius: 999px;
    padding: 8px 16px;
    font-family: var(--sans);
    font-size: 12px;
    font-weight: 600;
    cursor: pointer;
  }}
  .gaps-toggle:hover {{ border-color: #fff; color: #fff; }}
  .gaps-panel {{ display: none; background: #0a0a0a; border-top: 1px solid #1f1f1f; padding: 32px 24px; }}
  .gaps-panel.open {{ display: block; }}
  .gaps-panel-inner {{ max-width: var(--max); margin: 0 auto; color: #d4d2cc; font-family: var(--sans); font-size: 13px; }}
  .gaps-panel-title {{ font-family: var(--serif); font-size: 22px; color: #fff; margin-bottom: 6px; font-weight: 700; }}
  .gaps-panel-sub {{ font-size: 12px; color: #8d8a83; margin-bottom: 20px; }}
  .gap-section {{ margin-bottom: 22px; padding-bottom: 18px; border-bottom: 1px solid #1f1f1f; }}
  .gap-section:last-child {{ border-bottom: 0; }}
  .gap-heading {{ font-family: var(--sans); font-size: 11px; font-weight: 700; letter-spacing: 0.1em; text-transform: uppercase; color: #fff; margin-bottom: 12px; }}
  .gap-badge-thin {{ color: #ffb74a; font-size: 10px; margin-left: 6px; padding: 1px 6px; border: 1px solid #4a3920; border-radius: 2px; }}
  .gap-feed-row {{ padding: 8px 0; border-top: 1px dashed #2a2a2a; }}
  .gap-feed-row:first-of-type {{ border-top: 0; }}
  .gap-feed-name {{ font-weight: 600; color: #fff; }}
  .gap-feed-reason {{ color: #8d8a83; margin-left: 8px; font-style: italic; }}
  .gap-url-row {{ display: flex; gap: 8px; align-items: center; margin-top: 4px; }}
  .gap-feed-url {{ font-family: var(--mono); font-size: 11px; color: #8d8a83; flex: 1; overflow: hidden; text-overflow: ellipsis; }}
  .copy-btn {{ background: transparent; color: #d4d2cc; border: 1px solid #3a3a3a; border-radius: 4px; padding: 3px 8px; font-family: var(--sans); font-size: 11px; cursor: pointer; }}
  .copy-btn:hover {{ border-color: #fff; color: #fff; }}
  .gaps-empty {{ color: #8d8a83; font-style: italic; }}

  /* ─── Bookmark button (when image is hidden) ───────────────── */
  .story-card.no-thumb .bookmark-btn {{
    position: static;
    background: transparent;
    align-self: flex-end;
    width: 22px; height: 22px;
  }}
</style>
</head>
<body>

<!-- ── Top bar ──────────────────────────────────────────────────── -->
<div class="topbar">
  <div class="topbar-inner">
    <div class="topbar-left">
      <button class="icon-btn" aria-label="Menu" title="Sections">
        <svg width="16" height="12" viewBox="0 0 16 12" fill="none" stroke="currentColor" stroke-width="1.6"><line x1="0" y1="1" x2="16" y2="1"/><line x1="0" y1="6" x2="16" y2="6"/><line x1="0" y1="11" x2="16" y2="11"/></svg>
      </button>
      <span class="live-pill" id="live-pill">
        <span class="live-dot"></span>
        <span id="live-time">Updated {live_time_str}</span>
      </span>
    </div>
    <div class="wordmark">
      <span class="ed">ECON</span><span class="nd">Digest</span>
    </div>
    <div class="topbar-right">
      <div class="search-wrap" id="search-wrap">
        <button class="icon-btn" id="search-btn" aria-label="Search">
          <svg width="16" height="16" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.6"><circle cx="7" cy="7" r="5"/><line x1="11" y1="11" x2="15" y2="15"/></svg>
        </button>
        <input type="search" id="search" placeholder="Search stories…" autocomplete="off" aria-label="Search">
      </div>
      <button class="icon-btn" id="dark-btn" aria-label="Toggle dark mode" title="Toggle theme">
        <svg id="dark-svg" width="16" height="16" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.5"><path d="M13 9.5A5 5 0 1 1 6.5 3a4 4 0 0 0 6.5 6.5z"/></svg>
      </button>
      <button class="pill-btn ghost" onclick="window.location.href='archive/index.html'">Archive</button>
      <button class="pill-btn" id="topstories-cta">★ Top Stories</button>
    </div>
  </div>
</div>

<!-- ── Primary nav ──────────────────────────────────────────────── -->
<nav class="navbar" aria-label="Topic categories">
  <div class="nav-inner">
{nav_tabs_html}
  </div>
</nav>
<nav class="regionbar" aria-label="Filter by region">
  <div class="region-inner">
    <span class="region-label">Region</span>
    <button class="region-tab active" data-region="__all__">All</button>
    <button class="region-tab" data-region="ireland">Ireland</button>
    <button class="region-tab" data-region="uk">UK</button>
    <button class="region-tab" data-region="europe">Europe</button>
    <button class="region-tab" data-region="us">US</button>
    <button class="region-tab" data-region="canada">Canada</button>
    <button class="region-tab" data-region="latam">Latin America</button>
    <button class="region-tab" data-region="asia-pacific">Asia-Pacific</button>
    <button class="region-tab" data-region="africa-me">Africa &amp; Middle East</button>
    <button class="region-tab" data-region="global">Global</button>
    <button class="region-tab" data-region="underreported">⌕ Solo source</button>
  </div>
</nav>

<!-- ── Main ─────────────────────────────────────────────────────── -->
<main id="main-content">
  <noscript><p style="padding:2rem;color:#666">Enable JavaScript to view this digest.</p></noscript>
</main>
<div id="no-results">No stories match your search.</div>

<!-- ── Gaps panel (collapsible above footer) ───────────────────── -->
<div class="gaps-panel" id="gaps-panel">
  <div class="gaps-panel-inner" id="gaps-panel-inner"></div>
</div>

<!-- ── Footer ──────────────────────────────────────────────────── -->
<footer>
  <div class="footer-inner">
    <div>
      <div class="footer-brand">ECONDigest</div>
      <div class="footer-meta">
        Economic News Aggregator · {now_str}<br>
        {total} articles from the last {hours_back}h · Summaries by <strong>{model_str}</strong>
      </div>
    </div>
    <button class="gaps-toggle" id="gaps-toggle">📡 Feed Gaps (<span id="gaps-count">…</span>)</button>
  </div>
</footer>

<script>
// ── Embedded data ─────────────────────────────────────────────────────────
const DATA = {data_json};
const REGION_LABELS = {{
  'ireland':      'Ireland',
  'uk':           'UK',
  'europe':       'Europe',
  'us':           'US',
  'canada':       'Canada',
  'latam':        'Latin America',
  'asia-pacific': 'Asia-Pacific',
  'africa-me':    'Africa & ME',
  'global':       'Global',
}};

// ── Utilities ─────────────────────────────────────────────────────────────
function esc(s) {{
  return String(s ?? '')
    .replace(/&/g,'&amp;').replace(/</g,'&lt;')
    .replace(/>/g,'&gt;').replace(/"/g,'&quot;');
}}
function ageShort(publishedStr) {{
  if (!publishedStr) return '';
  const parts = publishedStr.split(', ');
  if (parts.length < 2) return '';
  const timeClean = parts[1].replace('UTC', '').trim();
  const dateStr = parts[0] + ' ' + timeClean + ':00 UTC';
  const pub = new Date(dateStr);
  if (isNaN(pub.getTime())) return '';
  const diffMs  = Date.now() - pub.getTime();
  const diffMin = Math.floor(diffMs / 60000);
  if (diffMin < 60) {{
    const m = Math.max(1, diffMin);
    return m + ' min ago';
  }}
  const diffHr = Math.floor(diffMin / 60);
  if (diffHr < 24) return diffHr + ' hr' + (diffHr === 1 ? '' : 's') + ' ago';
  const diffDays = Math.floor(diffHr / 24);
  return diffDays + ' day' + (diffDays === 1 ? '' : 's') + ' ago';
}}

// ── Saved bookmarks ───────────────────────────────────────────────────────
const SAVED_KEY = 'econ-saved';
function getSaved() {{
  try {{ return new Set(JSON.parse(localStorage.getItem(SAVED_KEY) || '[]')); }}
  catch {{ return new Set(); }}
}}
function persistSaved(set) {{ localStorage.setItem(SAVED_KEY, JSON.stringify([...set])); }}
function isSaved(url) {{ return getSaved().has(url); }}
function toggleSave(url) {{
  const saved = getSaved();
  if (saved.has(url)) saved.delete(url); else saved.add(url);
  persistSaved(saved); updateSavedTab();
}}
function updateSavedTab() {{
  const count = getSaved().size;
  const tab = document.getElementById('saved-tab');
  if (tab) tab.innerHTML = count > 0 ? '★ Saved <span class="tab-count">' + count + '</span>' : '★ Saved';
}}

// ── Story renderer (editorial card) ───────────────────────────────────────
function storyHTML(art, opts) {{
  opts = opts || {{}};
  const variant = opts.variant || 'standard';   // 'standard' | 'hero-center' | 'hero-flank' | 'compact'
  const showSummary = opts.showSummary !== false;
  const searchable = [art.title, art.source, art.summary, ...(art.kws || [])].join(' ').toLowerCase();
  const rslug  = art.region_slug || 'global';
  const rlabel = REGION_LABELS[rslug] || 'Global';
  const age    = ageShort(art.published);
  const saved  = isSaved(art.link);
  const hasImg = !!art.image;

  const thumb = hasImg
    ? '<a class="thumb-wrap" href="' + esc(art.link) + '" target="_blank" rel="noopener noreferrer"><div class="thumb" style="background-image:url(\\'' + esc(art.image) + '\\')"></div>' +
      '<button class="bookmark-btn' + (saved ? ' saved' : '') + '" data-url="' + esc(art.link) + '" aria-label="Save">' + (saved ? '★' : '☆') + '</button></a>'
    : '<div class="thumb-wrap"><div class="thumb no-image" data-label="' + esc((art.kws && art.kws[0]) || 'photo') + '"></div>' +
      '<button class="bookmark-btn' + (saved ? ' saved' : '') + '" data-url="' + esc(art.link) + '" aria-label="Save">' + (saved ? '★' : '☆') + '</button></div>';

  const flag = opts.isTop ? '<span class="top-flag">Top story</span>' : '';
  const headline = '<h3 class="headline"><a href="' + esc(art.link) + '" target="_blank" rel="noopener noreferrer">' + esc(art.title) + '</a></h3>';

  let teaser = '';
  if (showSummary && art.summary) {{
    const collapsedClass = opts.collapsed ? ' collapsed' : '';
    teaser = '<p class="teaser' + collapsedClass + '">' + esc(art.summary) + '</p>';
  }}

  // Meta line: SOURCE · AGE · [REGION] · solo?
  const metaParts = [];
  metaParts.push('<span class="meta-source">' + esc(art.source) + '</span>');
  if (age) metaParts.push('<span class="meta-sep">·</span><span>' + age + '</span>');
  metaParts.push('<span class="meta-sep">·</span><span class="meta-region">' + rlabel + '</span>');
  if (art.underreported) metaParts.push('<span class="solo-flag">Solo</span>');
  const meta = '<div class="meta-row">' + metaParts.join('') + '</div>';

  // Also-covering / perspectives → minimal text link
  let also = '';
  if (art.also && art.also.length) {{
    if (art.perspectives) {{
      const rows = art.also.map(a => {{
        const rs = a.region_slug || 'global';
        const rl = REGION_LABELS[rs] || 'Global';
        const snip = (a.summary || '').split(/(?<=[.!?])\\s+/)[0] || '';
        return '<div class="persp-row"><span class="meta-region">' + rl + '</span>' +
               '<a href="' + esc(a.url) + '" target="_blank" rel="noopener noreferrer">' + esc(a.source) + '</a>' +
               (snip ? '<span class="persp-snippet">' + esc(snip) + '</span>' : '') + '</div>';
      }}).join('');
      also = '<button class="also-link" data-toggle="persp">+ ' + (art.also.length + 1) + ' regional perspectives</button>' +
             '<div class="also-list" style="display:none">' + rows + '</div>';
    }} else {{
      const rows = art.also.map(a =>
        '<li><a href="' + esc(a.url) + '" target="_blank" rel="noopener noreferrer">' + esc(a.title) + '</a><span class="also-src"> — ' + esc(a.source) + '</span></li>'
      ).join('');
      also = '<button class="also-link" data-toggle="also">+ ' + art.also.length + ' more source' + (art.also.length > 1 ? 's' : '') + '</button>' +
             '<ul class="also-list" style="display:none">' + rows + '</ul>';
    }}
  }}

  const noThumbClass = (variant === 'hero-flank' && !hasImg) ? ' no-thumb' : '';
  const urAttr = art.underreported ? ' data-underreported="true"' : '';
  return '<div class="story-card ' + variant + noThumbClass + '" data-s="' + esc(searchable) + '" data-region="' + rslug + '"' + urAttr + '>' +
           thumb +
           flag +
           headline +
           teaser +
           meta +
           also +
         '</div>';
}}

function storyGridHTML(arts, opts) {{
  opts = opts || {{}};
  const cls = opts.cols === 3 ? 'three' : (opts.cols === 1 ? 'one' : '');
  return '<div class="story-grid ' + cls + '">' + arts.map(a => storyHTML(a, opts)).join('') + '</div>';
}}

// ── Hero (top stories layout) ─────────────────────────────────────────────
function heroHTML(stories) {{
  const left   = stories[1];
  const center = stories[0];
  const right  = stories[2];
  let html = '<section class="hero">';
  html += '<div class="hero-flank">' + (left ? storyHTML(left, {{ variant: 'hero-flank', isTop: true }}) : '') + '</div>';
  html += '<div class="hero-center">' + (center ? storyHTML(center, {{ variant: 'hero-center', isTop: true }}) : '') + '</div>';
  html += '<div class="hero-flank">' + (right ? storyHTML(right, {{ variant: 'hero-flank', isTop: true }}) : '') + '</div>';
  html += '</section>';
  return html;
}}

// ── Page views ────────────────────────────────────────────────────────────
function renderHome() {{
  const topKws = (DATA.topKeywords || []).slice(0, 10);
  const pillsHTML = topKws.map(kw =>
    '<button class="theme-pill" data-kw="' + esc(kw.keyword) + '">' + esc(kw.keyword) + '</button>'
  ).join('');
  const s = DATA.stats || {{}};
  const statsStr = (s.totalArticles || DATA.totalArticles) + ' articles · ' +
                   (s.feedsChecked || '?') + ' feeds · last ' + (s.hoursBack || DATA.hours) + 'h · ' + esc(s.model || '');

  let html = '<div class="page-head">' +
             '<h1 class="page-title">Economic Digest</h1>' +
             '<div class="page-title-sub">' + (DATA.generated || '') + ' — curated from ' + (s.feedsChecked || 0) + ' sources, scored for relevance</div>' +
             '</div>';

  // Hero
  if (DATA.topStories && DATA.topStories.length) {{
    html += heroHTML(DATA.topStories.slice(0, 3));
  }}

  // Themes bar
  html += '<div class="themes-bar">' +
            '<div class="themes-left">' +
              '<span class="themes-label">Today\\'s themes</span>' +
              '<div class="themes-pills">' + pillsHTML + '</div>' +
            '</div>' +
            '<div class="themes-right">' + statsStr + '</div>' +
          '</div>';

  // Categories
  html += '<div class="home-grid">';
  for (const cat of DATA.categories) {{
    const preview = cat.articles.slice(0, 4);
    const extra = cat.count - preview.length;
    html += '<section class="cat-section" data-cat="' + esc(cat.slug) + '">' +
              '<hr class="cat-rule">' +
              '<div class="cat-section-head">' +
                '<span class="cat-section-title">' + esc(cat.name) + '<span class="cnt">' + cat.count + ' stories</span></span>' +
                (extra > 0 ? '<button class="see-all-btn" data-cat="' + esc(cat.slug) + '">See all ' + cat.count + ' →</button>' : '') +
              '</div>' +
              storyGridHTML(preview, {{ collapsed: true, showSummary: true }}) +
            '</section>';
  }}
  html += '</div>';
  return html;
}}

function renderTopStories() {{
  let html = '<div class="page-head">' +
             '<h1 class="page-title">★ Top Stories <span class="page-title-count">' + DATA.topStories.length + '</span></h1>' +
             '<div class="page-title-sub">The day\\'s highest-scoring articles across all categories</div>' +
             '</div>';
  html += storyGridHTML(DATA.topStories, {{ cols: 3, isTop: true, showSummary: true }});
  return html;
}}

function renderCategory(cat) {{
  let html = '<button class="back-btn" id="back-btn">← All topics</button>' +
             '<div class="page-head">' +
             '<h1 class="page-title">' + esc(cat.name) + '<span class="page-title-count">' + cat.count + '</span></h1>' +
             '<div class="page-title-sub">' + cat.count + ' stories curated in the last 24 hours</div>' +
             '</div>';

  if (cat.subcats && cat.subcats.length > 1) {{
    for (const sub of cat.subcats) {{
      html += '<section class="subcat-section" data-subcat="' + esc(sub.slug) + '">' +
                '<h3 class="subcat-heading">' + esc(sub.name) + ' <span class="subcat-count">(' + sub.articles.length + ')</span></h3>' +
                storyGridHTML(sub.articles, {{ cols: 3, showSummary: true }}) +
              '</section>';
    }}
  }} else {{
    html += storyGridHTML(cat.articles, {{ cols: 3, showSummary: true }});
  }}
  return html;
}}

function renderSaved() {{
  const saved = getSaved();
  const seen = new Set(); const all = [];
  for (const cat of DATA.categories)
    for (const art of cat.articles)
      if (!seen.has(art.link)) {{ seen.add(art.link); all.push(art); }}
  for (const art of DATA.topStories)
    if (!seen.has(art.link)) {{ seen.add(art.link); all.push(art); }}
  const savedArts = all.filter(a => saved.has(a.link));

  let html = '<div class="page-head">' +
             '<h1 class="page-title">★ Saved <span class="page-title-count">' + savedArts.length + '</span></h1>' +
             '<div class="page-title-sub">Stories you marked for later. Stored in this browser only.</div>' +
             '</div>';
  if (savedArts.length === 0) {{
    html += '<div style="text-align:center;padding:80px 24px;color:var(--muted);font-family:var(--serif);font-size:18px;line-height:1.5;">' +
            'No saved stories yet.<br><span style="font-size:14px;">Tap the ☆ on any card to save it.</span></div>';
  }} else {{
    html += '<button id="clear-saved-btn" class="pill-btn ghost" style="margin-bottom:24px">Clear all saved</button>' +
            storyGridHTML(savedArts, {{ cols: 3, showSummary: true }});
  }}
  return html;
}}

// ── State + nav ───────────────────────────────────────────────────────────
let currentView = '__home__';
let searchQuery = '';
let regionFilter = '__all__';

function showView(catSlug) {{
  currentView = catSlug;
  const main = document.getElementById('main-content');

  document.querySelectorAll('.nav-tab').forEach(t =>
    t.classList.toggle('active', t.dataset.cat === catSlug));

  if (catSlug === '__home__') {{
    main.innerHTML = renderHome();
  }} else if (catSlug === '__top__') {{
    main.innerHTML = renderTopStories();
  }} else if (catSlug === '__saved__') {{
    main.innerHTML = renderSaved();
    const clearBtn = document.getElementById('clear-saved-btn');
    if (clearBtn) clearBtn.addEventListener('click', () => {{
      persistSaved(new Set()); updateSavedTab(); showView('__saved__');
    }});
  }} else {{
    const cat = DATA.categories.find(c => c.slug === catSlug);
    if (!cat) {{ main.innerHTML = '<p style="padding:2rem;color:var(--muted)">Category not found.</p>'; return; }}
    main.innerHTML = renderCategory(cat);
    const back = document.getElementById('back-btn');
    if (back) back.addEventListener('click', () => showView('__home__'));
  }}

  // Wire "see all"
  main.querySelectorAll('.see-all-btn').forEach(btn =>
    btn.addEventListener('click', () => showView(btn.dataset.cat)));

  // Bookmarks
  main.querySelectorAll('.bookmark-btn').forEach(btn =>
    btn.addEventListener('click', (e) => {{
      e.preventDefault(); e.stopPropagation();
      const url = btn.dataset.url;
      toggleSave(url);
      const nowSaved = isSaved(url);
      btn.textContent = nowSaved ? '★' : '☆';
      btn.classList.toggle('saved', nowSaved);
    }}));

  // Also-link expanders
  main.querySelectorAll('.also-link').forEach(btn =>
    btn.addEventListener('click', () => {{
      const list = btn.nextElementSibling;
      if (!list) return;
      const open = list.style.display === 'block';
      list.style.display = open ? 'none' : 'block';
      btn.style.color = open ? '' : 'var(--ink)';
    }}));

  // Read-summary toggles (collapsed teasers on home grid)
  main.querySelectorAll('.read-summary').forEach(btn =>
    btn.addEventListener('click', () => {{
      const card = btn.closest('.story-card');
      const t = card && card.querySelector('.teaser');
      if (t) t.classList.toggle('collapsed');
    }}));

  if (searchQuery || regionFilter !== '__all__') applyFilters();
  window.scrollTo(0, 0);
}}

// ── Filters ───────────────────────────────────────────────────────────────
function applyFilters() {{
  let visible = 0;
  document.querySelectorAll('.story-card').forEach(item => {{
    const matchSearch = !searchQuery || (item.dataset.s || '').includes(searchQuery);
    const matchRegion = regionFilter === '__all__' ||
      (regionFilter === 'underreported' ? item.dataset.underreported === 'true' : item.dataset.region === regionFilter);
    const show = matchSearch && matchRegion;
    item.classList.toggle('search-hidden', !show);
    if (show) visible++;
  }});
  ['subcat-section', 'cat-section'].forEach(cls => {{
    document.querySelectorAll('.' + cls).forEach(sec => {{
      sec.classList.toggle('search-hidden', !sec.querySelector('.story-card:not(.search-hidden)'));
    }});
  }});
  const active = searchQuery || regionFilter !== '__all__';
  document.getElementById('no-results').style.display = (active && visible === 0) ? 'block' : 'none';
}}

// ── Dark mode ─────────────────────────────────────────────────────────────
const DM_KEY = 'econ-dark';
function applyDark(on) {{
  document.body.classList.toggle('dark', on);
  const svg = document.getElementById('dark-svg');
  if (svg) svg.innerHTML = on
    ? '<circle cx="8" cy="8" r="3.2"/><line x1="8" y1="0.5" x2="8" y2="2.5"/><line x1="8" y1="13.5" x2="8" y2="15.5"/><line x1="0.5" y1="8" x2="2.5" y2="8"/><line x1="13.5" y1="8" x2="15.5" y2="8"/><line x1="2.7" y1="2.7" x2="4.1" y2="4.1"/><line x1="11.9" y1="11.9" x2="13.3" y2="13.3"/><line x1="2.7" y1="13.3" x2="4.1" y2="11.9"/><line x1="11.9" y1="4.1" x2="13.3" y2="2.7"/>'
    : '<path d="M13 9.5A5 5 0 1 1 6.5 3a4 4 0 0 0 6.5 6.5z"/>';
}}

// ── Init ──────────────────────────────────────────────────────────────────
(function () {{
  applyDark(!!localStorage.getItem(DM_KEY));

  document.getElementById('dark-btn').addEventListener('click', () => {{
    const on = !document.body.classList.contains('dark');
    localStorage.setItem(DM_KEY, on ? '1' : '');
    applyDark(on);
  }});

  // Search expand
  const sw = document.getElementById('search-wrap');
  const sb = document.getElementById('search-btn');
  const si = document.getElementById('search');
  sb.addEventListener('click', (e) => {{
    e.preventDefault();
    const open = sw.classList.toggle('open');
    if (open) si.focus();
    else {{ si.value = ''; searchQuery = ''; applyFilters(); }}
  }});
  si.addEventListener('input', function () {{
    searchQuery = this.value.toLowerCase().trim();
    applyFilters();
  }});
  si.addEventListener('keydown', (e) => {{ if (e.key === 'Escape') sb.click(); }});

  document.getElementById('topstories-cta').addEventListener('click', () => showView('__top__'));

  document.querySelectorAll('.nav-tab').forEach(tab =>
    tab.addEventListener('click', () => showView(tab.dataset.cat)));

  document.querySelectorAll('.region-tab').forEach(tab =>
    tab.addEventListener('click', () => {{
      regionFilter = tab.dataset.region;
      document.querySelectorAll('.region-tab').forEach(t => t.classList.toggle('active', t === tab));
      showView(currentView);
    }}));

  // Theme pill click → search
  document.addEventListener('click', function (e) {{
    const pill = e.target.closest('.theme-pill');
    if (!pill) return;
    const kw = pill.dataset.kw;
    const searchEl = document.getElementById('search');
    document.getElementById('search-wrap').classList.add('open');
    searchEl.value = kw;
    searchQuery = kw.toLowerCase().trim();
    applyFilters();
  }});

  updateSavedTab();
  showView('__home__');
  initGapsPanel();
}})();

// ── Gaps panel ────────────────────────────────────────────────────────────
function toggleGapsPanel() {{
  document.getElementById('gaps-panel').classList.toggle('open');
}}
function copyUrl(btn) {{
  const url = btn.dataset.url;
  navigator.clipboard.writeText(url).then(() => {{
    btn.textContent = '✓ Copied';
    setTimeout(() => {{ btn.textContent = '📋 Copy'; }}, 2000);
  }}).catch(() => {{}});
}}
function initGapsPanel() {{
  const gaps = DATA.feedGaps || [];
  const total = gaps.reduce((n, g) => n + (g.suggestions || []).length, 0);
  document.getElementById('gaps-count').textContent = total;
  document.getElementById('gaps-toggle').addEventListener('click', toggleGapsPanel);

  const inner = document.getElementById('gaps-panel-inner');
  const hasContent = gaps.some(g => (g.suggestions || []).length > 0);
  if (!gaps.length || !hasContent) {{
    inner.innerHTML = '<div class="gaps-panel-title">📡 Coverage Gaps & Suggested Feeds</div>' +
                      '<p class="gaps-empty" style="margin-top:8px">No significant coverage gaps identified today.</p>';
    return;
  }}
  let html = '<div class="gaps-panel-title">📡 Coverage Gaps & Suggested Feeds</div>' +
             '<div class="gaps-panel-sub">Identified after today\\'s run — review before adding to RSS feeds.</div>';
  for (const gap of gaps) {{
    const isThin = gap.article_count < 2;
    const badge = isThin ? ' <span class="gap-badge-thin">Thin</span>' : '';
    html += '<div class="gap-section">' +
              '<div class="gap-heading">' + esc(gap.category) + ' — ' + esc(gap.region) + ' (' + gap.article_count + ' article' + (gap.article_count === 1 ? '' : 's') + ' today)' + badge + '</div>';
    if (!gap.suggestions || !gap.suggestions.length) {{
      html += '<div class="gaps-empty">No suggestions available for this gap.</div>';
    }} else {{
      for (const s of gap.suggestions) {{
        const safeUrl = esc(s.url || '');
        html += '<div class="gap-feed-row">' +
                  '<span class="gap-feed-name">' + esc(s.name || '') + '</span>' +
                  (s.reason ? '<span class="gap-feed-reason">' + esc(s.reason) + '</span>' : '') +
                  '<div class="gap-url-row">' +
                    '<code class="gap-feed-url">' + safeUrl + '</code>' +
                    '<button class="copy-btn" data-url="' + safeUrl + '" onclick="copyUrl(this)">📋 Copy</button>' +
                  '</div>' +
                '</div>';
      }}
    }}
    html += '</div>';
  }}
  inner.innerHTML = html;
}}
</script>

</body>
</html>
"""


# =============================================================================
# MAIN ENTRY POINT
# =============================================================================

def main():
    print()
    print("=" * 62)
    print("  Economic News Aggregator")
    _now_utc = datetime.now(timezone.utc)
    _d = _now_utc.strftime('%d').lstrip('0') or '0'
    print(f"  {_now_utc.strftime(f'%A, {_d} %B %Y at %H:%M GMT')}")
    print("=" * 62)

    # ── 1. Load sources ───────────────────────────────────────────────────────
    print("\n[1/6] Reading RSS Feeds.docx...")
    sources        = read_sources()
    disabled_count = count_disabled_feeds()
    print(f"      Found {len(sources)} active feed sources ({disabled_count} disabled)")

    if not sources:
        sys.exit("No sources found in RSS Feeds.docx — check the file format.")

    # ── 2. Load keywords ──────────────────────────────────────────────────────
    print("\n[2/6] Reading Keywords.docx...")
    keywords = read_keywords()
    themes   = sorted(set(k['theme'] for k in keywords))
    print(f"      Found {len(keywords)} keywords across {len(themes)} themes:")
    for t in themes:
        count = sum(1 for k in keywords if k['theme'] == t)
        print(f"        • {t} ({count} keywords)")

    if not keywords:
        sys.exit("No keywords found in Keywords.docx — check the file format.")

    # ── 3. Fetch feeds ────────────────────────────────────────────────────────
    print(f"\n[3/6] Fetching all {len(sources)} RSS feeds...")
    raw_entries = fetch_all_feeds(sources)

    # ── 4. Filter to 24 hours ─────────────────────────────────────────────────
    print(f"\n[4/6] Filtering to articles from the last {HOURS_BACK} hours...")
    recent = filter_recent(raw_entries)

    if not recent:
        sys.exit(
            "\nNo recent articles found.\n"
            "Check your internet connection, or some feeds may not publish timestamps."
        )

    # ── 5. Score and deduplicate ──────────────────────────────────────────────
    print("\n[5/6] Scoring articles by keyword relevance and removing duplicates...")
    articles = build_articles(recent, keywords)
    articles = deduplicate(articles)

    if not articles:
        sys.exit(
            "\nNo relevant articles found after keyword filtering.\n"
            "Try broadening your keywords in Keywords.docx."
        )

    # Count per-feed contributions from scored, deduplicated articles
    feed_article_counts: dict = {}
    for art in articles:
        feed_article_counts[art['source']] = feed_article_counts.get(art['source'], 0) + 1

    print(f"\n      Top 10 articles by score:")
    for i, art in enumerate(articles[:10], 1):
        print(f"        {i:>2}. [Score {art['score']:>2}] {art['title'][:60]}")

    # ── 6. Generate summaries ─────────────────────────────────────────────────
    print(f"\n[6/6] Generating AI summaries via Ollama ({OLLAMA_MODEL})...")
    articles = add_summaries(articles)

    # ── 6.5. Cluster related stories ─────────────────────────────────────────
    print("\n[6.5/6] Clustering stories about the same event...")
    articles = cluster_articles(articles)

    ur_count = mark_underreported(articles)
    if ur_count:
        print(f"  Underreported: {ur_count} high-relevance solo-source article(s) flagged")

    # ── Write output ──────────────────────────────────────────────────────────
    print("\nComputing keyword hit frequencies...")
    kw_counts    = _count_all_keywords(articles, keywords)
    top_keywords = compute_top_keywords(kw_counts, keywords)
    if top_keywords:
        print(f"  Top themes: {', '.join(k['keyword'] for k in top_keywords[:5])} …")

    print("Updating keyword history...")
    history           = update_keyword_history(kw_counts)
    trending_keywords = compute_trending_keywords(kw_counts, history, keywords)
    if trending_keywords:
        print(f"  Trending: {', '.join(k['keyword'] for k in trending_keywords[:3])} …")
    else:
        print("  No spiking keywords (need ≥2 days of history for trends)")

    print("Identifying coverage gaps and querying Ollama for feed suggestions...")
    feed_gaps = suggest_missing_feeds(articles)

    prev_digest = _find_prev_digest()
    if prev_digest:
        print(f"\nPrevious digest found: {prev_digest}")

    print(f"\nWriting digest to: {OUTPUT_HTML}")
    digest = build_html(
        articles,
        prev_digest=prev_digest,
        top_keywords=top_keywords,
        feeds_checked=len(sources),
        trending_keywords=trending_keywords,
        feed_gaps=feed_gaps,
    )
    with open(OUTPUT_HTML, 'w', encoding='utf-8') as f:
        f.write(digest)

    print("\nArchiving digest...")
    archive_digest(digest, len(articles))

    print("Building trends page...")
    trends_html = build_trends_page(history, trending_keywords)
    trends_path = os.path.join(BASE_DIR, 'archive', 'trends.html')
    with open(trends_path, 'w', encoding='utf-8') as f:
        f.write(trends_html)
    print(f"  Trends page → {trends_path}")

    print("Building feed dashboard...")
    feeds_html = build_feed_dashboard(sources, feed_article_counts, disabled_count)
    feeds_path = os.path.join(BASE_DIR, 'archive', 'feeds.html')
    with open(feeds_path, 'w', encoding='utf-8') as f:
        f.write(feeds_html)
    print(f"  Feed dashboard → {feeds_path}")

    print()
    print("=" * 62)
    print(f"  Done!  {len(articles)} articles written to digest.html")
    print(f"  Open digest.html in any browser to read your digest.")
    print("=" * 62)
    print()

    # ── Email digest ──────────────────────────────────────────────────────────
    send_digest_email(OUTPUT_HTML)


def send_digest_email(digest_path):
    """Attach digest.html to an email and send via Gmail SMTP."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.base import MIMEBase
    from email.mime.text import MIMEText
    from email import encoders

    if not EMAIL_APP_PASSWORD:
        print("  Email skipped — add your Gmail App Password to EMAIL_APP_PASSWORD in aggregator.py")
        print("  Get one at: https://myaccount.google.com/apppasswords")
        return

    now = datetime.now(timezone.utc)
    day = now.strftime('%d').lstrip('0') or '0'
    subject = now.strftime(f"Economic News Digest — {day} %B %Y")

    msg = MIMEMultipart()
    msg['From']    = EMAIL_FROM
    msg['To']      = EMAIL_TO
    msg['Subject'] = subject

    msg.attach(MIMEText(
        f"Your daily Economic News Digest is attached.\n\n"
        f"Generated {now.strftime(f'{day} %B %Y at %H:%M GMT')}",
        'plain'
    ))

    filename = now.strftime(f"digest_{day}_%b_%Y.html").lower()
    with open(digest_path, 'rb') as f:
        part = MIMEBase('text', 'html')
        part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
        msg.attach(part)

    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(EMAIL_FROM, EMAIL_APP_PASSWORD)
            server.sendmail(EMAIL_FROM, EMAIL_TO, msg.as_string())
        print(f"  Digest emailed to {EMAIL_TO}")
    except Exception as e:
        print(f"  Email failed: {e}")


if __name__ == '__main__':
    main()
